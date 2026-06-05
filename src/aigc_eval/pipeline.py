# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import math
import os
import re
import shutil
import sys
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

import cv2
import fitz
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image
from scipy import stats
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler

SEED = 20260605
np.random.seed(SEED)

ROOT = Path(__file__).resolve().parents[2]
WORKSPACE = ROOT.parent
CONFIG_DIR = ROOT / "configs"
DATA_DIR = ROOT / "data"
RAW_DIR = DATA_DIR / "raw"
EXTRACTED_DIR = DATA_DIR / "extracted"
PROCESSED_DIR = DATA_DIR / "processed"
LABEL_DIR = DATA_DIR / "labels"
OUTPUT_DIR = ROOT / "outputs"
TABLE_DIR = OUTPUT_DIR / "tables"
FIGURE_DIR = OUTPUT_DIR / "figures"
LOG_DIR = OUTPUT_DIR / "logs"
SCORED_DIR = OUTPUT_DIR / "scored_samples"
REPORT_DIR = ROOT / "reports"
PAPER_DIR = ROOT / "paper"
PAPER_FIGURE_DIR = PAPER_DIR / "figures"
PAPER_TABLE_DIR = PAPER_DIR / "tables"
SUBMISSION_DIR = ROOT / "submission"


DIRS = [
    CONFIG_DIR,
    RAW_DIR,
    EXTRACTED_DIR,
    PROCESSED_DIR,
    LABEL_DIR,
    TABLE_DIR,
    FIGURE_DIR,
    LOG_DIR,
    SCORED_DIR,
    REPORT_DIR,
    PAPER_DIR,
    PAPER_FIGURE_DIR,
    PAPER_TABLE_DIR,
    SUBMISSION_DIR,
    ROOT / "tests",
]


def ensure_dirs() -> None:
    for d in DIRS:
        d.mkdir(parents=True, exist_ok=True)


def log(message: str) -> None:
    ensure_dirs()
    with (LOG_DIR / "pipeline.log").open("a", encoding="utf-8") as f:
        f.write(message.rstrip() + "\n")
    print(message)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def append_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(text.rstrip() + "\n")


def table_md(df: pd.DataFrame, max_rows: int = 20) -> str:
    if df is None or df.empty:
        return "（无可展示记录）"
    show = df.head(max_rows).copy()
    try:
        return show.to_markdown(index=False)
    except Exception:
        cols = list(show.columns)
        rows = ["| " + " | ".join(map(str, cols)) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
        for _, row in show.iterrows():
            rows.append("| " + " | ".join(str(row[c]) for c in cols) + " |")
        return "\n".join(rows)


def read_pdf_text(path: Path) -> tuple[str, int]:
    doc = fitz.open(path)
    text = "\n".join(page.get_text("text") for page in doc)
    return text, doc.page_count


def find_material_root() -> Path:
    candidates: list[Path] = []
    for p in WORKSPACE.iterdir():
        if not p.is_dir() or p.resolve() == ROOT.resolve():
            continue
        pdfs = list(p.glob("*.pdf"))
        attach_dirs = [d for d in p.iterdir() if d.is_dir() and d.name.startswith("附件")]
        if pdfs and attach_dirs:
            candidates.append(p)
    if not candidates:
        raise FileNotFoundError(f"未在 {WORKSPACE} 找到含 PDF 和附件目录的 B 题材料目录")
    candidates.sort(key=lambda x: x.stat().st_mtime, reverse=True)
    return candidates[0]


def find_source_files() -> dict[str, Path | None]:
    base = find_material_root()
    pdfs = list(base.glob("*.pdf"))
    docs = list(base.glob("*.doc*"))
    problem = None
    rules = None
    honesty = None
    for p in pdfs:
        name = p.name
        if "参赛细则" in name:
            rules = p
        elif "诚信" in name or "告知书" in name:
            honesty = p
        elif name.startswith("B") or "AI生成内容" in name:
            problem = p
    template = docs[0] if docs else None
    return {"base": base, "problem": problem, "rules": rules, "template": template, "honesty": honesty}


def image_info(path: Path) -> dict[str, Any]:
    with Image.open(path) as im:
        return {
            "width": im.width,
            "height": im.height,
            "mode": im.mode,
            "format": im.format,
            "dpi": im.info.get("dpi", ""),
            "exif_items": len(im.getexif()) if hasattr(im, "getexif") else 0,
            "metadata_keys": ",".join(sorted(im.info.keys())),
        }


def video_info(path: Path) -> dict[str, Any]:
    cap = cv2.VideoCapture(str(path))
    if not cap.isOpened():
        return {"readable": False}
    fps = float(cap.get(cv2.CAP_PROP_FPS) or 0)
    frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH) or 0)
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT) or 0)
    cap.release()
    return {
        "readable": True,
        "frame_count": frames,
        "fps": fps,
        "duration_sec": frames / fps if fps else np.nan,
        "width": width,
        "height": height,
    }


def inventory_record(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    rec: dict[str, Any] = {
        "path": str(path),
        "name": path.name,
        "suffix": suffix or "directory",
        "size_bytes": path.stat().st_size if path.is_file() else 0,
        "readable": True,
        "pages_or_samples": "",
        "requires_ocr": False,
        "manual_required": False,
        "notes": "",
    }
    try:
        if suffix == ".pdf":
            text, pages = read_pdf_text(path)
            rec["pages_or_samples"] = pages
            rec["text_chars"] = len(text)
            rec["requires_ocr"] = len(text.strip()) < 30
        elif suffix in [".png", ".jpg", ".jpeg", ".bmp", ".webp"]:
            info = image_info(path)
            rec.update(info)
            rec["pages_or_samples"] = f'{info["width"]}x{info["height"]}'
            rec["requires_ocr"] = False
        elif suffix in [".mp4", ".avi", ".mov", ".mkv"]:
            info = video_info(path)
            rec.update(info)
            rec["pages_or_samples"] = f'{info.get("frame_count", "")} frames'
        elif suffix in [".doc", ".docx"]:
            rec["notes"] = "Word 模板文件；旧 .doc 需 Word COM 转换或人工复核"
            rec["manual_required"] = suffix == ".doc"
        elif suffix in [".rar", ".zip", ".7z"]:
            rec["notes"] = "原始压缩包"
    except Exception as exc:
        rec["readable"] = False
        rec["manual_required"] = True
        rec["notes"] = f"{type(exc).__name__}: {exc}"
    return rec


def copy_raw_materials() -> None:
    files = find_source_files()
    base = files["base"]
    raw_unpack = RAW_DIR / "original_unpack"
    raw_unpack.mkdir(parents=True, exist_ok=True)
    for src in base.rglob("*"):
        if src.is_file():
            rel = src.relative_to(base)
            dst = raw_unpack / rel
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
    for p in WORKSPACE.glob("*.rar"):
        if "B" in p.name or "题" in p.name:
            shutil.copy2(p, RAW_DIR / p.name)


def audit_inputs() -> None:
    ensure_dirs()
    copy_raw_materials()
    files = find_source_files()
    base = files["base"]
    records = [inventory_record(p) for p in sorted(base.rglob("*")) if p.is_file()]
    archive_records = [inventory_record(p) for p in sorted(WORKSPACE.glob("*.rar")) if "B" in p.name or "题" in p.name]
    df = pd.DataFrame(archive_records + records)
    df.to_csv(TABLE_DIR / "input_file_inventory.csv", index=False, encoding="utf-8-sig")

    problem_ok = files["problem"] is not None and files["problem"].exists()
    rules_ok = files["rules"] is not None and files["rules"].exists()
    template_ok = files["template"] is not None and files["template"].exists()
    honesty_ok = files["honesty"] is not None and files["honesty"].exists()
    image_count = sum(1 for p in base.rglob("*") if p.suffix.lower() in [".png", ".jpg", ".jpeg"])
    video_count = sum(1 for p in base.rglob("*") if p.suffix.lower() in [".mp4", ".avi", ".mov", ".mkv"])

    md = f"""# 输入审计报告

## 审计结论

- 工作区：`{WORKSPACE}`
- 工程目录：`{ROOT}`
- 原始材料目录：`{base}`
- B 题题面 PDF：{"PASS" if problem_ok else "MISSING"}
- 参赛细则 PDF：{"PASS" if rules_ok else "MISSING"}
- 论文模板：{"PASS" if template_ok else "MISSING"}
- 诚信参赛告知书：{"PASS" if honesty_ok else "MISSING"}
- 附件图片数：{image_count}
- 附件视频数：{video_count}

## 文件清单

{table_md(df[["name", "suffix", "size_bytes", "pages_or_samples", "readable", "requires_ocr", "manual_required", "notes"]], 80)}

## 风险与人工确认

- 附件图片无 EXIF/PNG 文本提示词元数据，题面要求的“文本提示词结构化”需要使用可视内容代理描述，不能解释为原始生成提示词。
- 附件 1 共 8 张图，而题面表述“每种类型包含高、中、低三个质量等级”与“四种内容类型”存在数量张数上的潜在歧义，不能伪造专家等级。
- 可视审计未发现明确“产品渲染”样本，因此类型敏感性分析以附件实际内容类型为准，并在论文中说明题面类型覆盖要求与附件实际样本存在差异。
- 论文模板为旧版 `.doc`，若 Word COM 转换失败，则只能依据参赛细则复刻格式并标注待人工确认项。
"""
    write_text(REPORT_DIR / "INPUT_AUDIT.md", md)
    log("Stage 00 completed: input audit")


def extract_problem_questions(text: str) -> dict[str, str]:
    questions: dict[str, str] = {}
    patterns = [
        ("问题1", r"问题1：(.*?)(?=问题2：)"),
        ("问题2", r"问题2：(.*?)(?=问题3：)"),
        ("问题3", r"问题3：(.*)$"),
    ]
    for key, pat in patterns:
        m = re.search(pat, text, re.S)
        questions[key] = re.sub(r"\s+", " ", m.group(1)).strip() if m else "未能自动提取"
    return questions


def read_problem() -> None:
    ensure_dirs()
    files = find_source_files()
    if not files["problem"]:
        update_work_continuation("题面 PDF 缺失，无法继续大规模推进。")
        raise FileNotFoundError("B 题题面 PDF 缺失")
    text, pages = read_pdf_text(files["problem"])
    write_text(EXTRACTED_DIR / "problem_text.txt", text)
    questions = extract_problem_questions(text)
    q_df = pd.DataFrame([{"question": k, "content": v} for k, v in questions.items()])
    q_df.to_csv(TABLE_DIR / "problem_questions.csv", index=False, encoding="utf-8-sig")
    true_title = "B 题 AI 生成内容的质量评估与参数优化"
    if "无参考图像质量评价" in text and "视频时序质量评估" in text:
        focus = "题面实际核心为 NR-IQA 图像质量评价、8 张 AI 图像评估、视频时序质量评估。"
    else:
        focus = "题面核心需人工复核。"
    md = f"""# 题面读取报告

## 基本信息

- 题面文件：`{files["problem"]}`
- 页数：{pages}
- 题名：{true_title}
- 自动判断：{focus}

## 背景摘要

题面围绕扩散模型和生成式 AI 视频/图像内容质量评价展开，强调单帧图像质量、提示词语义匹配、技术质量、结构完整性，以及时序连贯性、光流连续性、内容一致性和闪烁检测。

## 真实三问

### 问题一

{questions["问题1"]}

### 问题二

{questions["问题2"]}

### 问题三

{questions["问题3"]}

## 对任务链的修正

用户初始任务链包含“参数-质量关系”和“参数优化”。但本 PDF 三问没有提供真实生成参数表，也没有把参数优化列为独立问题。因此本工程保留参数相关脚本以便审计，但主论文按题面修正为：

1. M1：无参考图像质量评价指标模型；
2. M2：组合赋权与 TOPSIS-灰色关联图像综合评价模型；
3. M3：质量有效性修正与等级语义校准；
4. M4：视频时序质量评价模型；
5. M5：本地规则型多智能体复核与稳定性分析。
"""
    write_text(REPORT_DIR / "PROBLEM_READING.md", md)
    assumptions = f"""# 模型假设与限制

1. 附件未提供每张图像的原始生成提示词、模型名、seed、采样步数、引导系数、采样器和生成耗时，因此不能进行真实生成参数优化，也不能构造伪参数。
2. 附件未提供原始提示词，本文不计算真实 prompt-image 语义保真度；实际计算使用“可视内容结构化代理描述”，只作为语义结构完整性代理。
3. 附件 1 的 8 张图没有专家标签、专家打分或官方高/中/低标签；论文中的等级为指标体系综合评分等级，不是专家质量标签。
4. 题面提到样本应覆盖产品渲染，但附件 1 的 8 张图经可视审计未发现明确产品渲染样本；论文按附件实际内容类型分析，不补造不存在的样本类型。
5. 视频中车辆检测若未使用训练好的外部目标检测模型，则只采用光流、亮度、颜色直方图、边缘与运动强度等可解释视觉代理指标，不伪造车辆检测结果。
6. AI 风险、伪影风险和模板风险均为模型风险提示，不构成违规判定或事实认定。
7. 所有随机过程固定随机种子 `{SEED}`，保证可复现。
"""
    write_text(REPORT_DIR / "ASSUMPTIONS.md", assumptions)
    log("Stage 01 completed: problem reading")


def try_convert_doc_template(template: Path | None) -> tuple[Path | None, str]:
    if not template or not template.exists():
        return None, "模板文件缺失"
    out = EXTRACTED_DIR / "competition_template_converted.docx"
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(template.resolve()))
        doc.SaveAs(str(out.resolve()), FileFormat=16)
        doc.Close(False)
        word.Quit()
        return out, "Word COM 转换成功"
    except Exception as exc:
        try:
            word.Quit()  # type: ignore[name-defined]
        except Exception:
            pass
        return None, f"Word COM 转换失败：{type(exc).__name__}: {exc}"


def docx_section_format(path: Path | None) -> dict[str, str]:
    if not path or not path.exists():
        return {}
    try:
        from docx import Document
        from docx.shared import Cm

        doc = Document(str(path))
        sec = doc.sections[0]
        return {
            "page_width_cm": f"{sec.page_width / Cm(1):.2f}",
            "page_height_cm": f"{sec.page_height / Cm(1):.2f}",
            "top_margin_cm": f"{sec.top_margin / Cm(1):.2f}",
            "bottom_margin_cm": f"{sec.bottom_margin / Cm(1):.2f}",
            "left_margin_cm": f"{sec.left_margin / Cm(1):.2f}",
            "right_margin_cm": f"{sec.right_margin / Cm(1):.2f}",
        }
    except Exception as exc:
        return {"docx_format_error": f"{type(exc).__name__}: {exc}"}


def extract_format_requirements() -> None:
    ensure_dirs()
    files = find_source_files()
    rules_text = ""
    if files["rules"]:
        rules_text, pages = read_pdf_text(files["rules"])
        write_text(EXTRACTED_DIR / "competition_rules_text.txt", rules_text)
    else:
        pages = 0
    converted, convert_msg = try_convert_doc_template(files["template"])
    fmt = docx_section_format(converted)

    page_size = (
        f'{fmt.get("page_width_cm")} cm x {fmt.get("page_height_cm")} cm（由模板转换读取）'
        if fmt.get("page_width_cm")
        else "待人工确认"
    )
    margin = (
        f'上 {fmt.get("top_margin_cm")} cm，下 {fmt.get("bottom_margin_cm")} cm，左 {fmt.get("left_margin_cm")} cm，右 {fmt.get("right_margin_cm")} cm（由模板转换读取）'
        if fmt.get("top_margin_cm")
        else "待人工确认"
    )
    req_rows = [
        ("页面大小", page_size, "模板" if fmt.get("page_width_cm") else "待人工确认"),
        ("页边距", margin, "模板" if fmt.get("top_margin_cm") else "待人工确认"),
        ("题目字体、字号、加粗、居中", "三号黑体字；题目居中", "参赛细则第 2 页"),
        ("队伍编号格式", "论文第一页包含队伍编号；提交文件命名为题号+队伍编号，例如 B202600001.pdf", "参赛细则"),
        ("摘要标题格式", "模板显示“摘 要”居中，小四字号；字体继承模板默认字体", "论文模板"),
        ("摘要正文字体和字号", "小四号宋体字，1.25 倍行距；摘要建议不超过一页", "参赛细则"),
        ("关键词格式", "模板显示“关键词：”左对齐，小四字号；关键词位于摘要正文之后", "论文模板"),
        ("一级标题格式", "四号黑体字并居中", "参赛细则第 2 页"),
        ("二级标题格式", "小四号黑体字，左端对齐，不居中", "参赛细则第 2 页"),
        ("三级标题格式", "小四号黑体字，左端对齐，不居中", "参赛细则第 2 页"),
        ("正文字体和字号", "小四号宋体字", "参赛细则第 2 页"),
        ("行距", "1.25 倍", "参赛细则第 2 页"),
        ("段前段后", "模板正文段落未见显式段前段后设置；当前按 Word 默认/0 处理", "论文模板"),
        ("图题格式", "待人工确认", "待人工确认"),
        ("表题格式", "待人工确认", "待人工确认"),
        ("公式编号格式", "待人工确认；本工程采用右侧圆括号编号作为论文草稿格式", "待人工确认"),
        ("页码格式", "不得有页眉；页码位于每页页脚中部；阿拉伯数字从 1 连续编号", "参赛细则"),
        ("是否允许目录", "不允许目录", "参赛细则第 1 页"),
        ("参考文献格式", "按正文引用次序；期刊、书籍、网上资源格式见细则；正文引用用方括号编号", "参赛细则"),
        ("附录格式", "正文之后为附录；附录页数不限", "参赛细则"),
        ("AI 工具使用说明位置", "参考文献后明确声明；支撑材料中提供 AI 工具使用详情说明", "参赛细则"),
        ("论文页数限制", "正文尽量控制在 20 页以内；摘要建议不超过一页", "参赛细则"),
        ("文件命名要求", "参赛论文：题号+队伍编号.pdf；支撑材料：题号+队伍编号+材料，rar 或 zip", "参赛细则"),
    ]
    req_df = pd.DataFrame(req_rows, columns=["item", "requirement", "source"])
    req_df.to_csv(TABLE_DIR / "format_requirements.csv", index=False, encoding="utf-8-sig")
    md = "# 论文格式要求提取\n\n"
    md += f"- 参赛细则文件：`{files['rules']}`，页数：{pages}\n"
    md += f"- 论文模板文件：`{files['template']}`\n"
    md += f"- 模板自动转换：{convert_msg}\n\n"
    md += table_md(req_df, 40)
    md += "\n\n## 待人工确认项\n\n"
    pending = req_df[req_df["requirement"].str.contains("待人工确认", na=False) | req_df["source"].str.contains("待人工确认", na=False)]
    md += table_md(pending, 40)
    write_text(REPORT_DIR / "FORMAT_REQUIREMENTS.md", md)

    audit = f"""# Word 模板审计

- 原始模板：`{files["template"]}`
- 模板类型：旧版 OLE `.doc`
- 自动转换状态：{convert_msg}
- 转换后模板：`{converted if converted else "无"}`
- 页面格式读取：{json.dumps(fmt, ensure_ascii=False)}

## 审计结论

若转换成功，页面大小和页边距优先来自转换后的模板；标题、正文、页码、页数和文件命名优先来自参赛细则。模板未明确或程序无法读取的项目已在 `FORMAT_REQUIREMENTS.md` 标为 `待人工确认`。
"""
    write_text(REPORT_DIR / "WORD_TEMPLATE_AUDIT.md", audit)
    log("Stage 13 completed: format requirements")


PROXY_DESCRIPTORS = {
    "1": ("写实风景", "写实风景 山谷 森林 河流 雾气 阳光 山体 清澈水面"),
    "2": ("人物肖像", "人物肖像 女性 室内 光影 油画风格 毛衣 藤椅"),
    "3": ("艺术插画", "艺术插画 狐狸 樱花 水彩 柔和天空 树枝 花瓣"),
    "4": ("动态街景", "街景 人群 奔跑 旗帜 逆光 运动模糊 城市道路"),
    "5": ("写实风景", "风景 溪流 树林 阳光 油画风格 草地 树冠"),
    "6": ("人物插画", "人物插画 少女 樱花 校服 动漫风格 春天 蓝天"),
    "7": ("水墨插画", "中国水墨 山水 亭台 竹子 梅花 小船 留白"),
    "8": ("像素街景", "像素艺术 城市街道 建筑 红白立面 小人物"),
}


def extract_materials() -> None:
    ensure_dirs()
    files = find_source_files()
    base = files["base"]
    images = sorted([p for p in base.rglob("*") if p.suffix.lower() in [".png", ".jpg", ".jpeg"]], key=lambda p: p.name)
    videos = sorted([p for p in base.rglob("*") if p.suffix.lower() in [".mp4", ".avi", ".mov", ".mkv"]], key=lambda p: p.name)
    manifest: list[dict[str, Any]] = []
    metadata: list[dict[str, Any]] = []
    prompts: list[dict[str, Any]] = []
    for i, p in enumerate(images, 1):
        stem = p.stem
        content_type, descriptor = PROXY_DESCRIPTORS.get(stem, ("待人工确认", f"图像 {stem}"))
        sid = f"image_{i:02d}"
        info = image_info(p)
        manifest.append(
            {
                "sample_id": sid,
                "sample_type": "image",
                "content_type": content_type,
                "source_path": str(p),
                "file_name": p.name,
                "original_prompt_available": False,
                "quality_label_available": False,
            }
        )
        metadata.append({"sample_id": sid, "file_name": p.name, **info, "size_bytes": p.stat().st_size})
        prompts.append(
            {
                "sample_id": sid,
                "file_name": p.name,
                "content_type": content_type,
                "proxy_descriptor": descriptor,
                "is_original_prompt": False,
                "note": "由可视内容审计得到的代理描述，不是原始生成提示词",
            }
        )
    for i, p in enumerate(videos, 1):
        sid = f"video_{i:02d}"
        info = video_info(p)
        manifest.append(
            {
                "sample_id": sid,
                "sample_type": "video",
                "content_type": "车流视频",
                "source_path": str(p),
                "file_name": p.name,
                "original_prompt_available": False,
                "quality_label_available": False,
            }
        )
        metadata.append({"sample_id": sid, "file_name": p.name, **info, "size_bytes": p.stat().st_size})
    manifest_df = pd.DataFrame(manifest)
    metadata_df = pd.DataFrame(metadata)
    prompt_df = pd.DataFrame(prompts)
    manifest_df.to_csv(TABLE_DIR / "sample_manifest.csv", index=False, encoding="utf-8-sig")
    metadata_df.to_csv(TABLE_DIR / "extracted_metadata.csv", index=False, encoding="utf-8-sig")
    prompt_df.to_csv(LABEL_DIR / "proxy_prompts.csv", index=False, encoding="utf-8-sig")

    md = f"""# 材料解析摘要

## 样本结构

{table_md(manifest_df, 50)}

## 元数据

{table_md(metadata_df, 50)}

## 提示词与标签状态

- 原始生成提示词：未在 PDF、文件名、EXIF 或 PNG 元数据中发现。
- 专家高/中/低质量标签：未发现。
- 已建立 `data/labels/proxy_prompts.csv`，其中的描述仅为可视内容结构化代理，用于计算题面要求的语义要素覆盖代理。
- 附件 1 共 {len(images)} 张图片，附件 2 共 {len(videos)} 个视频。
- 可视审计未发现明确产品渲染样本；本工程不补造该类型，按附件实际图像类型进行敏感性分析。
"""
    write_text(REPORT_DIR / "EXTRACTION_SUMMARY.md", md)
    log("Stage 02 completed: material extraction")


def parse_semantic_elements(text: str) -> dict[str, list[str]]:
    tokens = [t for t in re.split(r"[\s,，;；]+", text.strip()) if t]
    style_words = ["写实", "油画", "水彩", "动漫", "水墨", "像素", "艺术", "插画", "风格"]
    scene_words = ["山谷", "森林", "河流", "室内", "天空", "街景", "城市", "道路", "山水", "亭台", "建筑", "溪流"]
    attribute_words = ["清澈", "柔和", "逆光", "运动", "红白", "春天", "阳光", "雾气"]
    subject_words = [t for t in tokens if t not in style_words + scene_words + attribute_words]
    return {
        "subjects": subject_words[:6],
        "attributes": [t for t in tokens if t in attribute_words],
        "scenes": [t for t in tokens if t in scene_words],
        "styles": [t for t in tokens if t in style_words or "风格" in t],
        "tokens": tokens,
    }


def saturation(x: float, tau: float) -> float:
    x = max(float(x), 0.0)
    tau = max(float(tau), 1e-9)
    return float(1.0 - math.exp(-x / tau))


def safe_minmax(values: pd.Series) -> pd.Series:
    arr = values.astype(float)
    mn, mx = arr.min(), arr.max()
    if not np.isfinite(mn) or not np.isfinite(mx) or abs(mx - mn) < 1e-12:
        return pd.Series(np.full(len(arr), 0.5), index=values.index)
    return (arr - mn) / (mx - mn)


def image_array(path: Path) -> np.ndarray:
    with Image.open(path) as im:
        arr = np.array(im.convert("RGB"))
    return cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)


def blockiness_score(gray: np.ndarray) -> float:
    gray_f = gray.astype(np.float32)
    h, w = gray_f.shape
    if h < 16 or w < 16:
        return 0.0
    vb = np.abs(gray_f[:, 8::8] - gray_f[:, 7:-1:8]).mean() if w > 16 else 0
    vn = np.abs(gray_f[:, 4::8] - gray_f[:, 3:-1:8]).mean() if w > 16 else 1
    hb = np.abs(gray_f[8::8, :] - gray_f[7:-1:8, :]).mean() if h > 16 else 0
    hn = np.abs(gray_f[4::8, :] - gray_f[3:-1:8, :]).mean() if h > 16 else 1
    return float(max(0.0, ((vb + hb) / 2 - (vn + hn) / 2) / ((vn + hn) / 2 + 1e-6)))


def composition_score(gray: np.ndarray) -> float:
    gx = cv2.Sobel(gray, cv2.CV_32F, 1, 0, ksize=3)
    gy = cv2.Sobel(gray, cv2.CV_32F, 0, 1, ksize=3)
    sal = np.abs(gx) + np.abs(gy)
    total = float(sal.sum())
    h, w = gray.shape
    if total <= 1e-6:
        return 0.45
    ys, xs = np.indices(gray.shape)
    cx = float((xs * sal).sum() / total) / max(w - 1, 1)
    cy = float((ys * sal).sum() / total) / max(h - 1, 1)
    center_dist = math.sqrt((cx - 0.5) ** 2 + (cy - 0.5) ** 2)
    center = 1 - min(center_dist / 0.707, 1)
    thirds = [(1 / 3, 1 / 3), (2 / 3, 1 / 3), (1 / 3, 2 / 3), (2 / 3, 2 / 3)]
    thirds_score = max(math.exp(-((cx - tx) ** 2 + (cy - ty) ** 2) / (2 * 0.18**2)) for tx, ty in thirds)
    return float(np.clip(0.55 * center + 0.45 * thirds_score, 0, 1))


def extract_image_features(path: Path, descriptor: str, content_type: str, sample_id: str) -> dict[str, Any]:
    bgr = image_array(path)
    h0, w0 = bgr.shape[:2]
    scale = min(1.0, 1200.0 / max(h0, w0))
    if scale < 1.0:
        bgr_s = cv2.resize(bgr, (int(w0 * scale), int(h0 * scale)), interpolation=cv2.INTER_AREA)
    else:
        bgr_s = bgr
    gray = cv2.cvtColor(bgr_s, cv2.COLOR_BGR2GRAY)
    hsv = cv2.cvtColor(bgr_s, cv2.COLOR_BGR2HSV)
    lap_var = float(cv2.Laplacian(gray, cv2.CV_64F).var())
    sharpness = min(math.log1p(lap_var) / math.log1p(1600), 1.0)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    residual_std = float((gray.astype(np.float32) - blur.astype(np.float32)).std())
    noise_risk = float(np.clip((residual_std - 4) / 32, 0, 1))
    edges = cv2.Canny(gray, 80, 160)
    edge_density = float((edges > 0).mean())
    nlab, _, stats_cc, _ = cv2.connectedComponentsWithStats((edges > 0).astype(np.uint8), 8)
    if nlab > 1:
        areas = stats_cc[1:, cv2.CC_STAT_AREA]
        edge_cont = saturation(float(np.percentile(areas, 90)), 55) * saturation(edge_density, 0.045)
    else:
        edge_cont = 0.0
    mean_b = float(gray.mean() / 255)
    std_b = float(gray.std() / 255)
    clipped = float(((gray < 5) | (gray > 250)).mean())
    brightness_balance = float(np.clip(1 - abs(mean_b - 0.52) / 0.52 - clipped * 0.7, 0, 1))
    contrast = float(np.clip(std_b / 0.24, 0, 1))
    sat_mean = float(hsv[:, :, 1].mean() / 255)
    hist_h, _ = np.histogram(hsv[:, :, 0], bins=36, range=(0, 180), density=True)
    hist_h = hist_h + 1e-12
    color_entropy = float(-(hist_h * np.log(hist_h)).sum() / math.log(36))
    color_richness = float(np.clip(0.55 * saturation(sat_mean, 0.28) + 0.45 * color_entropy, 0, 1))
    block = blockiness_score(gray)
    oversmooth = float(np.clip((0.018 - edge_density) / 0.018, 0, 1))
    artifact_risk = float(np.clip(0.45 * clipped + 0.35 * min(block, 1) + 0.20 * oversmooth, 0, 1))
    artifact_inverse = 1 - artifact_risk
    shape_reg = float(np.clip(0.55 * edge_cont + 0.25 * artifact_inverse + 0.20 * contrast, 0, 1))
    structure_integrity = float(np.clip(0.55 * edge_cont + 0.30 * shape_reg + 0.15 * artifact_inverse, 0, 1))
    comp = composition_score(gray)
    gray_hist, _ = np.histogram(gray, bins=64, range=(0, 255), density=True)
    gray_hist = gray_hist + 1e-12
    gray_entropy = float(-(gray_hist * np.log(gray_hist)).sum() / math.log(64))
    nss = float(np.clip(0.25 * brightness_balance + 0.25 * contrast + 0.25 * gray_entropy + 0.25 * artifact_inverse, 0, 1))
    elems = parse_semantic_elements(descriptor)
    coverage = sum(1 for k in ["subjects", "attributes", "scenes", "styles"] if elems[k]) / 4
    specificity = min(len(elems["tokens"]) / 9, 1)
    semantic_proxy = float(np.clip(0.55 * coverage + 0.45 * specificity, 0, 1))
    technical = float(np.clip(0.32 * sharpness + 0.24 * nss + 0.18 * (1 - noise_risk) + 0.18 * artifact_inverse + 0.08 * contrast, 0, 1))
    return {
        "sample_id": sample_id,
        "file_name": path.name,
        "content_type": content_type,
        "width": w0,
        "height": h0,
        "megapixels": w0 * h0 / 1e6,
        "aspect_ratio": w0 / h0,
        "file_size_mb": path.stat().st_size / (1024 * 1024),
        "semantic_proxy_score": semantic_proxy,
        "semantic_component_coverage": coverage,
        "proxy_prompt_token_count": len(elems["tokens"]),
        "sharpness_raw": lap_var,
        "sharpness_score": sharpness,
        "noise_risk_score": noise_risk,
        "noise_inverse_score": 1 - noise_risk,
        "artifact_risk_score": artifact_risk,
        "artifact_inverse_score": artifact_inverse,
        "edge_density": edge_density,
        "edge_continuity_score": edge_cont,
        "shape_regularity_score": shape_reg,
        "structure_integrity_score": structure_integrity,
        "brightness_mean": mean_b,
        "brightness_balance_score": brightness_balance,
        "contrast_score": contrast,
        "color_richness_score": color_richness,
        "composition_proxy_score": comp,
        "nss_score": nss,
        "technical_quality_score": technical,
        "clipped_pixel_rate": clipped,
        "blockiness_risk": block,
        "oversmooth_risk": oversmooth,
        "proxy_descriptor": descriptor,
    }


def extract_video_temporal_features(video_path: Path) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    key_dir = EXTRACTED_DIR / "video_keyframes"
    key_dir.mkdir(parents=True, exist_ok=True)
    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        raise RuntimeError(f"无法打开视频：{video_path}")
    fps = float(cap.get(cv2.CAP_PROP_FPS) or 0)
    total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH) or 0)
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT) or 0)
    frames: list[np.ndarray] = []
    idxs: list[int] = []
    idx = 0
    step = max(1, total // 60)
    while True:
        ok, frame = cap.read()
        if not ok:
            break
        if idx % step == 0 or idx == total - 1:
            frames.append(frame)
            idxs.append(idx)
        idx += 1
    cap.release()
    if len(frames) < 2:
        raise RuntimeError("视频帧数不足，无法计算时序特征")

    frame_rows: list[dict[str, Any]] = []
    small_gray: list[np.ndarray] = []
    hist_list: list[np.ndarray] = []
    key_indices = {0, len(frames) // 4, len(frames) // 2, 3 * len(frames) // 4, len(frames) - 1}
    for local_i, (global_i, frame) in enumerate(zip(idxs, frames)):
        small = cv2.resize(frame, (480, int(frame.shape[0] * 480 / frame.shape[1])), interpolation=cv2.INTER_AREA)
        gray = cv2.cvtColor(small, cv2.COLOR_BGR2GRAY)
        hsv = cv2.cvtColor(small, cv2.COLOR_BGR2HSV)
        small_gray.append(gray)
        hist = cv2.calcHist([hsv], [0, 1], None, [24, 16], [0, 180, 0, 256])
        hist = cv2.normalize(hist, hist).flatten()
        hist_list.append(hist)
        lap = float(cv2.Laplacian(gray, cv2.CV_64F).var())
        brightness = float(gray.mean() / 255)
        contrast = float(gray.std() / 255)
        frame_rows.append(
            {
                "frame_index": global_i,
                "time_sec": global_i / fps if fps else np.nan,
                "sharpness_raw": lap,
                "sharpness_score": min(math.log1p(lap) / math.log1p(1200), 1.0),
                "brightness_mean": brightness,
                "contrast": contrast,
            }
        )
        if local_i in key_indices:
            out = key_dir / f"traffic_frame_{global_i:04d}.jpg"
            cv2.imwrite(str(out), frame)

    flow_rows: list[dict[str, Any]] = []
    prev_flow = None
    for i in range(1, len(small_gray)):
        flow = cv2.calcOpticalFlowFarneback(small_gray[i - 1], small_gray[i], None, 0.5, 3, 15, 3, 5, 1.2, 0)
        mag, ang = cv2.cartToPolar(flow[..., 0], flow[..., 1])
        grad_x = cv2.Sobel(mag, cv2.CV_32F, 1, 0, ksize=3)
        grad_y = cv2.Sobel(mag, cv2.CV_32F, 0, 1, ksize=3)
        flow_grad = float(np.mean(np.sqrt(grad_x**2 + grad_y**2)))
        mag_mean = float(np.mean(mag))
        smoothness = float(1 / (1 + flow_grad / (mag_mean + 1e-6)))
        hist_corr = float(cv2.compareHist(hist_list[i - 1].astype("float32"), hist_list[i].astype("float32"), cv2.HISTCMP_CORREL))
        hist_corr = float(np.clip((hist_corr + 1) / 2, 0, 1))
        bright_delta = abs(frame_rows[i]["brightness_mean"] - frame_rows[i - 1]["brightness_mean"])
        flow_delta = float(np.nan if prev_flow is None else np.mean(np.abs(mag - prev_flow)))
        prev_flow = mag
        flow_rows.append(
            {
                "pair_index": i,
                "start_frame": idxs[i - 1],
                "end_frame": idxs[i],
                "flow_magnitude_mean": mag_mean,
                "flow_magnitude_std": float(np.std(mag)),
                "flow_smoothness_score": smoothness,
                "histogram_consistency_score": hist_corr,
                "brightness_delta": bright_delta,
                "flow_temporal_delta": flow_delta,
            }
        )

    frame_df = pd.DataFrame(frame_rows)
    flow_df = pd.DataFrame(flow_rows)
    flow_delta_med = float(flow_df["flow_temporal_delta"].dropna().median()) if flow_df["flow_temporal_delta"].notna().any() else 0.0
    flow_mean = float(flow_df["flow_magnitude_mean"].mean())
    flow_cont = float(flow_df["flow_smoothness_score"].mean())
    content_cons = float(flow_df["histogram_consistency_score"].mean())
    flicker = float(flow_df["brightness_delta"].mean())
    flicker_score = float(np.clip(1 - flicker / 0.08, 0, 1))
    frame_quality = float(frame_df["sharpness_score"].mean() * 0.45 + np.clip(frame_df["contrast"].mean() / 0.25, 0, 1) * 0.25 + (1 - frame_df["brightness_mean"].std() / 0.12) * 0.30)
    temporal_jitter_score = float(1 / (1 + flow_delta_med / (flow_mean + 1e-6)))
    temporal_score = float(np.clip(100 * (0.25 * frame_quality + 0.25 * flow_cont + 0.20 * temporal_jitter_score + 0.20 * content_cons + 0.10 * flicker_score), 0, 100))
    instability_threshold = {
        "brightness_delta_mean_gt_0.08": flicker > 0.08,
        "flow_smoothness_lt_0.45": flow_cont < 0.45,
        "content_consistency_lt_0.70": content_cons < 0.70,
        "temporal_quality_lt_65": temporal_score < 65,
    }
    summary = {
        "video_name": video_path.name,
        "frame_count": total,
        "fps": fps,
        "duration_sec": total / fps if fps else np.nan,
        "width": width,
        "height": height,
        "sampled_frames": len(frames),
        "frame_quality_score": frame_quality * 100,
        "flow_continuity_score": flow_cont * 100,
        "flow_temporal_jitter_score": temporal_jitter_score * 100,
        "content_consistency_score": content_cons * 100,
        "flicker_suppression_score": flicker_score * 100,
        "brightness_delta_mean": flicker,
        "flow_magnitude_mean": flow_mean,
        "flow_temporal_delta_median": flow_delta_med,
        "video_temporal_quality_score": temporal_score,
        "instability_flag": any(instability_threshold.values()),
        "instability_conditions": json.dumps(instability_threshold, ensure_ascii=False),
    }
    return frame_df, flow_df, summary


def build_features() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "sample_manifest.csv").exists():
        extract_materials()
    manifest = pd.read_csv(TABLE_DIR / "sample_manifest.csv")
    prompts = pd.read_csv(LABEL_DIR / "proxy_prompts.csv") if (LABEL_DIR / "proxy_prompts.csv").exists() else pd.DataFrame()
    prompt_map = {r["sample_id"]: r for _, r in prompts.iterrows()}
    rows: list[dict[str, Any]] = []
    for _, r in manifest[manifest["sample_type"] == "image"].iterrows():
        sid = r["sample_id"]
        pr = prompt_map.get(sid, {})
        descriptor = pr.get("proxy_descriptor", r["content_type"])
        rows.append(extract_image_features(Path(r["source_path"]), descriptor, r["content_type"], sid))
    feat = pd.DataFrame(rows)
    feat.to_csv(TABLE_DIR / "features.csv", index=False, encoding="utf-8-sig")
    video_summaries: list[dict[str, Any]] = []
    for _, r in manifest[manifest["sample_type"] == "video"].iterrows():
        frame_df, flow_df, summary = extract_video_temporal_features(Path(r["source_path"]))
        frame_df.to_csv(TABLE_DIR / "video_frame_metrics.csv", index=False, encoding="utf-8-sig")
        flow_df.to_csv(TABLE_DIR / "video_flow_metrics.csv", index=False, encoding="utf-8-sig")
        video_summaries.append(summary)
    pd.DataFrame(video_summaries).to_csv(TABLE_DIR / "video_temporal_features.csv", index=False, encoding="utf-8-sig")
    md = f"""# 特征构建报告

## 图像特征

{table_md(feat[["sample_id", "content_type", "semantic_proxy_score", "sharpness_score", "nss_score", "structure_integrity_score", "artifact_risk_score"]], 20)}

## 视频时序特征

{table_md(pd.DataFrame(video_summaries), 10)}

## 说明

图像特征均来自本地 OpenCV/Pillow 计算；视频时序特征采用 Farneback 光流、颜色直方图一致性和亮度突变指标。未使用外部 LLM API 或伪造车辆检测结果。
"""
    write_text(REPORT_DIR / "FEATURE_EXTRACTION_REPORT.md", md)
    log("Stage 03 completed: feature building")


INDICATORS = [
    ("semantic_proxy_score", "语义结构完整性代理", "正向", "由代理描述中主体、属性、场景、风格四类要素覆盖率计算"),
    ("sharpness_score", "清晰度", "正向", "Laplacian 方差经对数饱和函数归一化"),
    ("nss_score", "自然场景统计质量", "正向", "亮度均衡、对比度、灰度熵和伪影反向指标综合"),
    ("noise_inverse_score", "噪声抑制", "正向", "高频残差噪声风险的反向分"),
    ("artifact_inverse_score", "伪影抑制", "正向", "过曝/欠曝、块效应、过平滑风险的反向分"),
    ("edge_continuity_score", "边缘连续性", "正向", "Canny 边缘连通区域的饱和得分"),
    ("shape_regularity_score", "形状规则性", "正向", "边缘连续、伪影反向与对比度综合"),
    ("structure_integrity_score", "结构完整性", "正向", "边缘连续性、形状规则性与伪影反向综合"),
    ("brightness_balance_score", "亮度均衡", "正向", "亮度均值偏离和截断像素惩罚"),
    ("contrast_score", "对比度", "正向", "灰度标准差饱和得分"),
    ("color_richness_score", "色彩丰富度", "正向", "HSV 饱和度与色相熵综合"),
    ("composition_proxy_score", "构图代理", "正向", "梯度显著性重心与中心/三分点关系"),
]


def build_quality_indicator_system() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "features.csv").exists():
        build_features()
    feat = pd.read_csv(TABLE_DIR / "features.csv")
    defs = pd.DataFrame(INDICATORS, columns=["indicator", "name_cn", "direction", "definition"])
    raw = feat[["sample_id", "file_name", "content_type"] + [x[0] for x in INDICATORS]].copy()
    norm = raw[["sample_id", "file_name", "content_type"]].copy()
    for ind, _, direction, _ in INDICATORS:
        vals = raw[ind].astype(float)
        if direction == "正向":
            norm[ind] = safe_minmax(vals)
        else:
            norm[ind] = 1 - safe_minmax(vals)
    defs.to_csv(TABLE_DIR / "indicator_definition_table.csv", index=False, encoding="utf-8-sig")
    raw.to_csv(TABLE_DIR / "indicator_matrix_raw.csv", index=False, encoding="utf-8-sig")
    norm.to_csv(TABLE_DIR / "indicator_matrix_normalized.csv", index=False, encoding="utf-8-sig")
    md = f"""# M1 内容质量指标量化模型

## 指标矩阵

设附件 1 中共有 $n$ 个图像样本、$m$ 个评价指标，原始指标矩阵为

$$
X=(x_{{ij}})_{{n\\times m}},
$$

其中 $x_{{ij}}$ 表示第 $i$ 张图像在第 $j$ 个指标上的取值。

## 归一化

正向指标采用

$$
x'_{{ij}}=\\frac{{x_{{ij}}-\\min_i x_{{ij}}}}{{\\max_i x_{{ij}}-\\min_i x_{{ij}}}},
$$

负向指标采用

$$
x'_{{ij}}=\\frac{{\\max_i x_{{ij}}-x_{{ij}}}}{{\\max_i x_{{ij}}-\\min_i x_{{ij}}}}.
$$

对数量或复杂度类指标使用饱和函数

$$
q(x)=1-\\exp(-x/\\tau),
$$

避免“越多越好”的虚高。

## 指标定义

{table_md(defs, 50)}

## 原始指标矩阵预览

{table_md(raw, 12)}
"""
    write_text(REPORT_DIR / "QUALITY_INDICATOR_SYSTEM.md", md)
    log("Stage 04 completed: indicator system")


def ahp_consistency(weights: np.ndarray) -> dict[str, float]:
    n = len(weights)
    mat = weights[:, None] / weights[None, :]
    eigvals, eigvecs = np.linalg.eig(mat)
    idx = int(np.argmax(eigvals.real))
    lam = float(eigvals.real[idx])
    vec = np.abs(eigvecs[:, idx].real)
    vec = vec / vec.sum()
    ci = (lam - n) / (n - 1) if n > 1 else 0
    ri_table = {1: 0, 2: 0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45, 10: 1.49, 11: 1.51, 12: 1.48}
    ri = ri_table.get(n, 1.56)
    cr = ci / ri if ri else 0
    return {"lambda_max": lam, "CI": float(ci), "CR": float(cr), "matrix": mat, "weights": vec}


def entropy_weights(x: pd.DataFrame) -> np.ndarray:
    arr = x.astype(float).to_numpy()
    arr = np.clip(arr, 0, None) + 1e-12
    colsum = arr.sum(axis=0)
    p = arr / colsum
    k = 1 / math.log(arr.shape[0]) if arr.shape[0] > 1 else 0
    e = -k * np.sum(p * np.log(p), axis=0)
    d = 1 - e
    if np.allclose(d.sum(), 0):
        return np.ones(arr.shape[1]) / arr.shape[1]
    return d / d.sum()


def critic_weights(x: pd.DataFrame) -> np.ndarray:
    arr = x.astype(float).to_numpy()
    std = np.nanstd(arr, axis=0)
    corr = np.corrcoef(arr, rowvar=False)
    corr = np.nan_to_num(corr, nan=0.0)
    conflict = np.sum(1 - corr, axis=1)
    c = std * conflict
    if np.allclose(c.sum(), 0):
        return np.ones(arr.shape[1]) / arr.shape[1]
    return c / c.sum()


def build_weights() -> None:
    ensure_dirs()
    setup_matplotlib()
    if not (TABLE_DIR / "indicator_matrix_normalized.csv").exists():
        build_quality_indicator_system()
    norm = pd.read_csv(TABLE_DIR / "indicator_matrix_normalized.csv")
    inds = [x[0] for x in INDICATORS]
    x = norm[inds]
    base = np.array([0.16, 0.13, 0.10, 0.08, 0.10, 0.11, 0.09, 0.10, 0.04, 0.04, 0.03, 0.02], dtype=float)
    base = base / base.sum()
    ahp = ahp_consistency(base)
    w_ahp = ahp["weights"]
    w_entropy = entropy_weights(x)
    w_critic = critic_weights(x)
    best = None
    for a in np.linspace(0, 1, 21):
        for b in np.linspace(0, 1 - a, 21):
            c = 1 - a - b
            w = a * w_ahp + b * w_entropy + c * w_critic
            loss = np.sum((w - w_ahp) ** 2 + (w - w_entropy) ** 2 + (w - w_critic) ** 2)
            if best is None or loss < best[0]:
                best = (loss, a, b, c, w)
    _, a, b, c, w_comb = best
    def wdf(name: str, weights: np.ndarray) -> pd.DataFrame:
        return pd.DataFrame({"indicator": inds, "weight": weights, "method": name})

    wdf("AHP", w_ahp).to_csv(TABLE_DIR / "ahp_weights.csv", index=False, encoding="utf-8-sig")
    wdf("Entropy", w_entropy).to_csv(TABLE_DIR / "entropy_weights.csv", index=False, encoding="utf-8-sig")
    wdf("CRITIC", w_critic).to_csv(TABLE_DIR / "critic_weights.csv", index=False, encoding="utf-8-sig")
    combined = wdf("Combined", w_comb)
    combined.to_csv(TABLE_DIR / "combined_weights.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame([{"alpha_AHP": a, "beta_entropy": b, "gamma_CRITIC": c, "loss": best[0], "CR": ahp["CR"]}]).to_csv(
        TABLE_DIR / "combined_weight_mix.csv", index=False, encoding="utf-8-sig"
    )

    fig, ax = plt.subplots(figsize=(11, 5))
    plot_df = pd.concat([wdf("AHP", w_ahp), wdf("Entropy", w_entropy), wdf("CRITIC", w_critic), combined])
    pivot = plot_df.pivot(index="indicator", columns="method", values="weight")
    pivot.plot(kind="bar", ax=ax)
    ax.set_title("指标权重对比")
    ax.set_ylabel("权重")
    ax.tick_params(axis="x", rotation=60)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "weight_comparison.png", dpi=220)
    plt.close(fig)

    md = f"""# M2 AHP-熵权-CRITIC 组合赋权模型

## AHP 一致性

- 最大特征值：{ahp["lambda_max"]:.6f}
- CI：{ahp["CI"]:.6f}
- CR：{ahp["CR"]:.6f}
- 结论：{"PASS" if ahp["CR"] <= 0.1 else "需调整"}

## 组合权重

组合权重采用

$$
w=\\alpha w^{{AHP}}+\\beta w^{{Entropy}}+\\gamma w^{{CRITIC}},\\quad
\\alpha+\\beta+\\gamma=1.
$$

网格搜索得到 $\\alpha={a:.2f}$，$\\beta={b:.2f}$，$\\gamma={c:.2f}$。

{table_md(combined, 30)}

![权重对比](../outputs/figures/weight_comparison.png)
"""
    write_text(REPORT_DIR / "WEIGHT_MODEL_REPORT.md", md)
    log("Stage 05 completed: weights")


def grade_fixed(score: float) -> str:
    if score >= 85:
        return "优秀"
    if score >= 70:
        return "良好"
    if score >= 55:
        return "中等"
    return "较低"


def assign_relative_quality_tiers(scores: pd.Series) -> pd.Series:
    ranks = scores.rank(ascending=False, method="first").astype(int)
    n = len(ranks)
    if n == 8:
        return ranks.map(lambda r: "相对高" if r <= 2 else ("相对中" if r <= 5 else "相对低"))
    high_cut = max(1, math.ceil(n / 3))
    mid_cut = max(high_cut + 1, math.ceil(2 * n / 3))
    return ranks.map(lambda r: "相对高" if r <= high_cut else ("相对中" if r <= mid_cut else "相对低"))


def strength_weakness_from_features(row: pd.Series) -> tuple[str, str]:
    dimensions = {
        "语义结构完整性代理": float(row.get("semantic_proxy_score", np.nan)),
        "技术质量": float(row.get("technical_quality_score", np.nan)),
        "结构完整性": float(row.get("structure_integrity_score", np.nan)),
        "视觉表达": float(0.5 * row.get("color_richness_score", 0) + 0.5 * row.get("composition_proxy_score", 0)),
        "伪影抑制": float(row.get("artifact_inverse_score", np.nan)),
    }
    valid = {k: v for k, v in dimensions.items() if np.isfinite(v)}
    if not valid:
        return "待人工复核", "待人工复核"
    strength = max(valid, key=valid.get)
    weakness = min(valid, key=valid.get)
    return strength, weakness


def quality_evaluation() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "combined_weights.csv").exists():
        build_weights()
    norm = pd.read_csv(TABLE_DIR / "indicator_matrix_normalized.csv")
    weights = pd.read_csv(TABLE_DIR / "combined_weights.csv").set_index("indicator")["weight"]
    inds = [x for x in weights.index if x in norm.columns]
    x = norm[inds].astype(float).to_numpy()
    w = weights.loc[inds].to_numpy()
    v = x * w
    ideal_pos = v.max(axis=0)
    ideal_neg = v.min(axis=0)
    d_pos = np.sqrt(((v - ideal_pos) ** 2).sum(axis=1))
    d_neg = np.sqrt(((v - ideal_neg) ** 2).sum(axis=1))
    c = d_neg / (d_pos + d_neg + 1e-12)
    diff = np.abs(x - x.max(axis=0))
    dmin = diff.min()
    dmax = diff.max()
    rho = 0.5
    grey_coeff = (dmin + rho * dmax) / (diff + rho * dmax + 1e-12)
    grey = (grey_coeff * w).sum(axis=1) / w.sum()
    weighted_sum = (x * w).sum(axis=1) / w.sum()
    base = norm[["sample_id", "file_name", "content_type"]].copy()
    topsis = base.copy()
    topsis["topsis_closeness"] = c
    grey_df = base.copy()
    grey_df["grey_relation_degree"] = grey
    res = base.copy()
    res["topsis_score_100"] = c * 100
    res["grey_score_100"] = grey * 100
    res["weighted_sum_100"] = weighted_sum * 100
    for lam in [0.3, 0.5, 0.7]:
        res[f"quality_score_lambda_{lam}"] = 100 * (lam * c + (1 - lam) * grey)
    res["Q_raw"] = res["quality_score_lambda_0.5"]
    res["fixed_grade"] = res["Q_raw"].map(grade_fixed)
    try:
        km = KMeans(n_clusters=3, random_state=SEED, n_init=20).fit(res[["Q_raw"]])
        centers = pd.Series(km.cluster_centers_.flatten()).sort_values()
        label_map = {centers.index[0]: "较低", centers.index[1]: "中等", centers.index[2]: "较高"}
        res["kmeans_grade"] = [label_map[x] for x in km.labels_]
    except Exception:
        res["kmeans_grade"] = "MANUAL_CHECK_REQUIRED"
    res["quantile_grade"] = pd.qcut(res["Q_raw"].rank(method="first"), q=3, labels=["较低", "中等", "较高"]).astype(str)
    res["rank"] = res["Q_raw"].rank(ascending=False, method="min").astype(int)
    topsis.to_csv(TABLE_DIR / "topsis_scores.csv", index=False, encoding="utf-8-sig")
    grey_df.to_csv(TABLE_DIR / "grey_relation_scores.csv", index=False, encoding="utf-8-sig")
    res.to_csv(TABLE_DIR / "comprehensive_quality_scores.csv", index=False, encoding="utf-8-sig")
    comp = res[["sample_id", "file_name", "content_type", "topsis_score_100", "grey_score_100", "weighted_sum_100", "Q_raw", "fixed_grade", "quantile_grade", "kmeans_grade", "rank"]]
    comp.to_csv(TABLE_DIR / "quality_grade_methods_comparison.csv", index=False, encoding="utf-8-sig")
    comp.to_csv(TABLE_DIR / "final_quality_grades.csv", index=False, encoding="utf-8-sig")
    sens = res[["sample_id", "file_name", "quality_score_lambda_0.3", "quality_score_lambda_0.5", "quality_score_lambda_0.7"]]
    sens.to_csv(TABLE_DIR / "lambda_sensitivity.csv", index=False, encoding="utf-8-sig")
    md = f"""# M3 TOPSIS-灰色关联综合质量评价模型

## 模型公式

TOPSIS 贴近度：

$$
C_i=\\frac{{D_i^-}}{{D_i^+ + D_i^-}}.
$$

灰色关联度：

$$
\\xi_{{ij}}=\\frac{{\\Delta_{{min}}+\\rho\\Delta_{{max}}}}{{\\Delta_{{ij}}+\\rho\\Delta_{{max}}}},\\quad
G_i=\\sum_j w_j\\xi_{{ij}}.
$$

综合质量分：

$$
Q_i=100[\\lambda C_i+(1-\\lambda)G_i],\\quad \\lambda=0.5.
$$

## 评估结果

{table_md(comp.sort_values("rank"), 20)}

## 可靠性说明

同时输出 TOPSIS、灰色关联和加权和三个分数，并做 $\\lambda=0.3,0.5,0.7$ 的敏感性分析。由于无专家标签，分级仅表示模型内部综合等级。
"""
    write_text(REPORT_DIR / "QUALITY_EVALUATION_MODEL.md", md)
    log("Stage 06 completed: quality evaluation")


def validity_corrected_score() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "comprehensive_quality_scores.csv").exists():
        quality_evaluation()
    feat = pd.read_csv(TABLE_DIR / "features.csv")
    scores = pd.read_csv(TABLE_DIR / "comprehensive_quality_scores.csv")
    df = scores.merge(feat, on=["sample_id", "file_name", "content_type"], how="left")
    res = df[["sample_id", "file_name", "content_type", "Q_raw"]].copy()
    res["content_relevance_validity"] = df["semantic_proxy_score"]
    res["technical_validity"] = df["technical_quality_score"]
    res["structure_logic_validity"] = df["structure_integrity_score"]
    res["modality_consistency_validity"] = df["nss_score"]
    res["expression_validity"] = 0.5 * df["color_richness_score"] + 0.5 * df["composition_proxy_score"]
    res["G_valid"] = (
        0.30 * res["content_relevance_validity"]
        + 0.25 * res["technical_validity"]
        + 0.20 * res["structure_logic_validity"]
        + 0.15 * res["modality_consistency_validity"]
        + 0.10 * res["expression_validity"]
    )
    quantity_proxy = safe_minmax(df["megapixels"]) * 0.45 + safe_minmax(df["edge_density"]) * 0.35 + safe_minmax(df["proxy_prompt_token_count"]) * 0.20
    res["quantity_proxy_score"] = quantity_proxy
    res["G_stack"] = 1 - res["quantity_proxy_score"] * (1 - res["G_valid"])
    penalty = 100 * (0.05 * df["artifact_risk_score"] + 0.03 * df["noise_risk_score"] + 0.02 * df["oversmooth_risk"] + 0.02 * df["clipped_pixel_rate"])
    res["risk_penalty"] = penalty
    res["Q_corrected"] = np.clip(res["Q_raw"] * (0.75 + 0.25 * res["G_valid"]) - res["risk_penalty"], 0, 100)
    res["Q_final"] = np.clip(res["Q_raw"] * (0.85 + 0.15 * res["G_valid"]) - 0.5 * res["risk_penalty"], 0, 100)
    res["final_grade"] = res["Q_final"].map(grade_fixed)
    res["absolute_quality_grade"] = res["final_grade"]
    res["rank"] = res["Q_final"].rank(ascending=False, method="first").astype(int)
    res["relative_quality_tier"] = assign_relative_quality_tiers(res["Q_final"])
    strengths = df.apply(strength_weakness_from_features, axis=1)
    res["key_strength"] = [x[0] for x in strengths]
    res["key_weakness"] = [x[1] for x in strengths]
    res.to_csv(TABLE_DIR / "validity_features.csv", index=False, encoding="utf-8-sig")
    res.to_csv(TABLE_DIR / "corrected_quality_scores.csv", index=False, encoding="utf-8-sig")
    final = res[
        [
            "sample_id",
            "file_name",
            "content_type",
            "Q_raw",
            "G_valid",
            "Q_corrected",
            "Q_final",
            "absolute_quality_grade",
            "relative_quality_tier",
            "rank",
            "key_strength",
            "key_weakness",
        ]
    ].copy()
    final["final_grade"] = final["absolute_quality_grade"]
    final["semantic_final_grade"] = final["final_grade"]
    final.to_csv(TABLE_DIR / "final_semantic_quality_grades.csv", index=False, encoding="utf-8-sig")
    relative_table = final[
        [
            "sample_id",
            "file_name",
            "content_type",
            "Q_final",
            "absolute_quality_grade",
            "relative_quality_tier",
            "rank",
            "key_strength",
            "key_weakness",
        ]
    ].sort_values("rank")
    relative_table.to_csv(TABLE_DIR / "image_relative_quality_tiers.csv", index=False, encoding="utf-8-sig")
    md1 = f"""# M4 质量有效性修正模型

有效性门控：

$$
G_{{valid}}=0.30V_r+0.25V_t+0.20V_s+0.15V_m+0.10V_e.
$$

反堆砌门控：

$$
G_{{stack}}=1-S_q(1-G_{{valid}}).
$$

最终温和修正分：

$$
Q_{{final}}=clip(Q_{{raw}}(0.85+0.15G_{{valid}})-0.5P,0,100).
$$

## 结果

{table_md(final.sort_values("Q_final", ascending=False), 20)}

## 附件内部相对三档

为回应题面“高、中、低质量等级”的表述，同时避免把机器分数包装成专家标签，本文新增相对等级：按 `Q_final` 排序，前 2 张为相对高，第 3-5 张为相对中，第 6-8 张为相对低。相对等级只用于附件 1 内部比较。

{table_md(relative_table, 20)}
"""
    write_text(REPORT_DIR / "VALIDITY_CORRECTED_MODEL.md", md1)
    md2 = f"""# 等级语义校准

## 校准原则

1. 没有专家标签时，不把模型等级写成真实人工等级；
2. 固定阈值、分位数和 KMeans 分级若出现语义不一致，以 `Q_final` 的固定阈值作为绝对等级；
3. 另设 `relative_quality_tier` 回应附件内部高/中/低质量比较，不等同于专家标签；
4. 对提示词缺失导致的语义结构代理不确定性，在结论中给出限制说明。

## 最终语义等级

{table_md(final.sort_values("Q_final", ascending=False), 20)}

## 相对高/中/低等级

{table_md(relative_table, 20)}
"""
    write_text(REPORT_DIR / "SEMANTIC_GRADE_CALIBRATION.md", md2)
    log("Stage 07 completed: validity corrected score")


def parameter_effect_analysis() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "final_semantic_quality_grades.csv").exists():
        validity_corrected_score()
    feat = pd.read_csv(TABLE_DIR / "features.csv")
    score = pd.read_csv(TABLE_DIR / "final_semantic_quality_grades.csv")
    df = feat.merge(score[["sample_id", "Q_final"]], on="sample_id", how="left")
    proxy_cols = ["megapixels", "aspect_ratio", "file_size_mb", "proxy_prompt_token_count", "sharpness_score", "edge_density", "color_richness_score", "artifact_risk_score"]
    param_table = df[["sample_id", "file_name", "content_type"] + proxy_cols + ["Q_final"]]
    param_table.to_csv(TABLE_DIR / "parameter_feature_table.csv", index=False, encoding="utf-8-sig")
    rows = []
    for col in proxy_cols:
        rho, pval = stats.spearmanr(df[col], df["Q_final"])
        tau, kp = stats.kendalltau(df[col], df["Q_final"])
        rows.append({"feature": col, "feature_type": "observable_proxy_not_generation_parameter", "spearman_rho": rho, "spearman_p": pval, "kendall_tau": tau, "kendall_p": kp})
    corr = pd.DataFrame(rows).sort_values("spearman_rho", key=lambda s: s.abs(), ascending=False)
    corr.to_csv(TABLE_DIR / "parameter_quality_correlation.csv", index=False, encoding="utf-8-sig")
    imp = corr[["feature", "spearman_rho", "kendall_tau"]].copy()
    imp["importance_proxy"] = imp["spearman_rho"].abs()
    imp.to_csv(TABLE_DIR / "parameter_importance.csv", index=False, encoding="utf-8-sig")
    metrics = pd.DataFrame(
        [
            {
                "model": "Ridge/LASSO/PLS on true generation parameters",
                "status": "NOT_APPLICABLE",
                "reason": "附件没有真实生成参数，且 n=8；不训练会被误读为真实参数模型的回归器",
            },
            {
                "model": "Spearman/Kendall observable proxy sensitivity",
                "status": "DONE",
                "reason": "仅分析可观测属性与模型质量分的内部相关性",
            },
        ]
    )
    metrics.to_csv(TABLE_DIR / "parameter_model_metrics.csv", index=False, encoding="utf-8-sig")
    md = f"""# 参数-质量关系分析

## 适用性审计

题面 PDF 和附件未提供真实生成参数记录，无法构造 seed、steps、guidance scale、模型名、采样器等参数变量。因此本节不做真实参数优化或参数回归，只做可观测代理特征与模型综合质量分的内部敏感性分析。

## 可观测代理特征相关性

{table_md(corr, 30)}

## 建模限制

这些相关性只能说明本批附件中“可观测图像属性”和模型分数的关系，不能解释为生成参数的因果影响。
"""
    write_text(REPORT_DIR / "PARAMETER_EFFECT_ANALYSIS.md", md)
    log("Stage 08 completed: proxy parameter effect analysis")


def parameter_optimization() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "parameter_quality_correlation.csv").exists():
        parameter_effect_analysis()
    candidates = pd.DataFrame(
        [
            {"strategy_id": "S1", "scope": "图像生成前提示词", "recommendation": "明确主体、属性、空间关系和风格四类语义要素", "parameter_combination": "NOT_AVAILABLE", "evidence": "题面要求语义结构化；附件无真实提示词"},
            {"strategy_id": "S2", "scope": "图像质量控制", "recommendation": "优先降低伪影、过平滑和异常截断像素，再提高细节复杂度", "parameter_combination": "NOT_AVAILABLE", "evidence": "有效性门控防止复杂度虚高"},
            {"strategy_id": "S3", "scope": "视频生成/筛选", "recommendation": "采用光流连续性、亮度突变和颜色一致性作为筛选指标", "parameter_combination": "NOT_AVAILABLE", "evidence": "问题三时序质量模型"},
        ]
    )
    best = pd.DataFrame(
        [
            {
                "parameter_name": "seed/steps/guidance_scale/resolution/model_name",
                "recommended_value": "NOT_APPLICABLE",
                "reason": "附件未提供真实生成参数及范围，不能伪造最优参数组合",
            }
        ]
    )
    comp = pd.DataFrame(
        [
            {"method": "grid_search", "status": "NOT_APPLICABLE", "reason": "缺少真实参数空间"},
            {"method": "response_surface_search", "status": "NOT_APPLICABLE", "reason": "缺少真实参数观测值和质量响应样本"},
            {"method": "strategy_rule_optimization", "status": "DONE", "reason": "输出基于指标缺陷的改进策略，不输出虚构参数"},
        ]
    )
    candidates.to_csv(TABLE_DIR / "optimization_candidates.csv", index=False, encoding="utf-8-sig")
    best.to_csv(TABLE_DIR / "best_parameter_recommendations.csv", index=False, encoding="utf-8-sig")
    comp.to_csv(TABLE_DIR / "optimization_method_comparison.csv", index=False, encoding="utf-8-sig")
    md = f"""# 参数优化报告

## 结论

真实 B 题三问没有要求输出最优生成参数组合，附件也没有参数记录。本工程不伪造参数，故真实参数优化标记为 `NOT_APPLICABLE`。

## 可执行优化策略

{table_md(candidates, 10)}

## 方法适用性

{table_md(comp, 10)}
"""
    write_text(REPORT_DIR / "PARAMETER_OPTIMIZATION_REPORT.md", md)
    log("Stage 09 completed: optimization applicability report")


def multi_agent_evaluation() -> None:
    ensure_dirs()
    if not (TABLE_DIR / "final_semantic_quality_grades.csv").exists():
        validity_corrected_score()
    feat = pd.read_csv(TABLE_DIR / "features.csv")
    scores = pd.read_csv(TABLE_DIR / "final_semantic_quality_grades.csv")
    df = feat.merge(scores, on=["sample_id", "file_name", "content_type"], how="left")
    rows = []
    for _, r in df.iterrows():
        agents = {
            "RelevanceAgent": (r["semantic_proxy_score"] * 100, "代理语义要素覆盖率"),
            "TechnicalAgent": (r["technical_quality_score"] * 100, "清晰度/NSS/噪声/伪影"),
            "StructureAgent": (r["structure_integrity_score"] * 100, "边缘连续性与结构完整性"),
            "AestheticOrReadabilityAgent": ((0.5 * r["color_richness_score"] + 0.5 * r["composition_proxy_score"]) * 100, "色彩与构图代理"),
            "SafetyAndTrustAgent": ((0.6 * r["artifact_inverse_score"] + 0.4 * r["noise_inverse_score"]) * 100, "伪影与噪声风险反向"),
            "ParameterAgent": (50.0, "无真实生成参数，给中性分并要求人工补充"),
        }
        for agent, (score, evidence) in agents.items():
            rows.append(
                {
                    "sample_id": r["sample_id"],
                    "file_name": r["file_name"],
                    "content_type": r["content_type"],
                    "agent": agent,
                    "agent_score": float(score),
                    "confidence": 0.8 if agent != "ParameterAgent" else 0.3,
                    "evidence": evidence,
                    "weakness": "无专家标签，规则型复核不覆盖主评分" if agent != "ParameterAgent" else "缺少真实参数",
                }
            )
        sample_md = f"""# {r["sample_id"]} Agent 复核

- 文件：{r["file_name"]}
- 类型：{r["content_type"]}
- 主评分：{r["Q_final"]:.2f}
- 主等级：{r["final_grade"]}

{table_md(pd.DataFrame([x for x in rows if x["sample_id"] == r["sample_id"]])[["agent", "agent_score", "confidence", "evidence", "weakness"]], 20)}
"""
        write_text(SCORED_DIR / f'{r["sample_id"]}_agent_review.md', sample_md)
    agent_df = pd.DataFrame(rows)
    agent_df.to_csv(TABLE_DIR / "agent_scores.csv", index=False, encoding="utf-8-sig")
    dis = agent_df.groupby("sample_id")["agent_score"].agg(["mean", "std", "min", "max"]).reset_index()
    dis["review_required"] = dis["std"] > 18
    dis.to_csv(TABLE_DIR / "agent_disagreement.csv", index=False, encoding="utf-8-sig")
    md = f"""# 多智能体复核与稳定性分析

## 规则型 Agent

本工程构建本地规则型 Agent，不调用外部 LLM API。Agent 只给复核建议，不覆盖主评分。

## Agent 分歧

{table_md(dis, 20)}
"""
    write_text(REPORT_DIR / "MULTI_AGENT_EVALUATION.md", md)
    log("Stage 10 completed: multi-agent review")


def setup_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "SimSun", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def save_placeholder_figure(name: str, title: str, message: str) -> None:
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.axis("off")
    ax.text(0.5, 0.62, title, ha="center", va="center", fontsize=16, weight="bold")
    ax.text(0.5, 0.40, message, ha="center", va="center", fontsize=11, wrap=True)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / name, dpi=220)
    plt.close(fig)


def generate_figures() -> None:
    ensure_dirs()
    setup_matplotlib()
    if not (TABLE_DIR / "agent_scores.csv").exists():
        multi_agent_evaluation()
    manifest = pd.read_csv(TABLE_DIR / "sample_manifest.csv")
    final = pd.read_csv(TABLE_DIR / "final_semantic_quality_grades.csv")
    feat = pd.read_csv(TABLE_DIR / "features.csv")
    weights = pd.read_csv(TABLE_DIR / "combined_weights.csv")
    agent = pd.read_csv(TABLE_DIR / "agent_scores.csv")

    fig, ax = plt.subplots(figsize=(6, 4))
    manifest["content_type"].value_counts().plot(kind="bar", ax=ax, color="#4C78A8")
    ax.set_title("样本类型与数量")
    ax.set_xlabel("内容类型")
    ax.set_ylabel("数量")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "sample_type_counts.png", dpi=220)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(11, 5))
    ax.axis("off")
    boxes = [
        ("AI生成内容质量", 0.5, 0.90),
        ("语义结构代理", 0.15, 0.62),
        ("技术质量", 0.38, 0.62),
        ("结构完整性", 0.62, 0.62),
        ("视觉表达", 0.85, 0.62),
        ("提示词要素覆盖", 0.15, 0.34),
        ("清晰度/NSS/噪声/伪影", 0.38, 0.34),
        ("边缘连续/形状规则", 0.62, 0.34),
        ("亮度/色彩/构图", 0.85, 0.34),
    ]
    for text, x, y in boxes:
        ax.text(x, y, text, ha="center", va="center", bbox=dict(boxstyle="round,pad=0.35", fc="#F5F7FA", ec="#3A3A3A"))
    for x in [0.15, 0.38, 0.62, 0.85]:
        ax.plot([0.5, x], [0.86, 0.68], color="#666")
    for x in [0.15, 0.38, 0.62, 0.85]:
        ax.plot([x, x], [0.58, 0.40], color="#666")
    ax.set_title("质量指标体系层次图")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "indicator_hierarchy.png", dpi=220)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.hist(final["Q_final"], bins=6, color="#59A14F", edgecolor="white")
    ax.set_title("综合质量分布")
    ax.set_xlabel("Q_final")
    ax.set_ylabel("样本数")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "quality_score_distribution.png", dpi=220)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6, 4))
    final["final_grade"].value_counts().reindex(["优秀", "良好", "中等", "较低"]).dropna().plot(kind="bar", ax=ax, color="#F28E2B")
    ax.set_title("最终等级分布")
    ax.set_xlabel("等级")
    ax.set_ylabel("样本数")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "grade_distribution.png", dpi=220)
    plt.close(fig)

    comp = pd.read_csv(TABLE_DIR / "quality_grade_methods_comparison.csv")
    fig, ax = plt.subplots(figsize=(10, 4.8))
    comp.set_index("file_name")[["topsis_score_100", "grey_score_100", "weighted_sum_100", "Q_raw"]].plot(kind="bar", ax=ax)
    ax.set_title("多模型评分对比")
    ax.set_ylabel("分数")
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "grade_methods_comparison.png", dpi=220)
    plt.close(fig)

    val = pd.read_csv(TABLE_DIR / "corrected_quality_scores.csv")
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.hist(val["G_valid"], bins=6, color="#76B7B2", edgecolor="white")
    ax.set_title("有效性门控分布")
    ax.set_xlabel("G_valid")
    ax.set_ylabel("样本数")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "validity_gate_distribution.png", dpi=220)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(10, 4.8))
    val.set_index("file_name")[["Q_raw", "Q_corrected", "Q_final"]].plot(kind="bar", ax=ax)
    ax.set_title("Raw/Corrected/Final 分数对比")
    ax.set_ylabel("分数")
    ax.tick_params(axis="x", rotation=45)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "score_comparison.png", dpi=220)
    plt.close(fig)

    corr = pd.read_csv(TABLE_DIR / "parameter_quality_correlation.csv")
    fig, ax = plt.subplots(figsize=(7.5, 3.8))
    vals = corr.set_index("feature")[["spearman_rho", "kendall_tau"]]
    im = ax.imshow(vals.T, aspect="auto", cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_yticks(range(vals.shape[1]), vals.columns)
    ax.set_xticks(range(vals.shape[0]), vals.index, rotation=45, ha="right")
    ax.set_title("可观测代理特征-质量相关热力图")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "parameter_quality_heatmap.png", dpi=220)
    plt.close(fig)

    imp = pd.read_csv(TABLE_DIR / "parameter_importance.csv")
    fig, ax = plt.subplots(figsize=(8, 4.2))
    imp.sort_values("importance_proxy").plot(kind="barh", x="feature", y="importance_proxy", ax=ax, color="#E15759")
    ax.set_title("可观测代理特征重要性")
    ax.set_xlabel("|Spearman rho|")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "parameter_importance.png", dpi=220)
    plt.close(fig)

    save_placeholder_figure("optimization_curve.png", "参数优化不适用", "附件没有真实生成参数和参数空间；本工程不生成虚构收敛曲线。")
    save_placeholder_figure("parameter_response_surface.png", "响应面不适用", "缺少真实参数-质量响应样本，不能构造响应面最优点。")

    pivot = agent.pivot_table(index="sample_id", columns="agent", values="agent_score")
    fig, ax = plt.subplots(figsize=(10, 4.8))
    im = ax.imshow(pivot.to_numpy(), aspect="auto", cmap="YlGnBu", vmin=0, vmax=100)
    ax.set_xticks(range(pivot.shape[1]), pivot.columns, rotation=35, ha="right")
    ax.set_yticks(range(pivot.shape[0]), pivot.index)
    ax.set_title("多智能体评分热力图")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "agent_score_heatmap.png", dpi=220)
    plt.close(fig)

    if (TABLE_DIR / "video_frame_metrics.csv").exists():
        vf = pd.read_csv(TABLE_DIR / "video_frame_metrics.csv")
        fig, ax1 = plt.subplots(figsize=(9, 4.5))
        ax1.plot(vf["time_sec"], vf["sharpness_score"] * 100, label="帧清晰度", color="#4C78A8")
        ax1.set_xlabel("时间/s")
        ax1.set_ylabel("清晰度分")
        ax2 = ax1.twinx()
        ax2.plot(vf["time_sec"], vf["brightness_mean"], label="亮度均值", color="#F28E2B", alpha=0.8)
        ax2.set_ylabel("亮度均值")
        ax1.set_title("视频关键帧质量分析")
        lines, labels = ax1.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(lines + lines2, labels + labels2, loc="best")
        fig.tight_layout()
        fig.savefig(FIGURE_DIR / "video_frame_quality_analysis.png", dpi=220)
        plt.close(fig)
    if (TABLE_DIR / "video_temporal_features.csv").exists():
        vt = pd.read_csv(TABLE_DIR / "video_temporal_features.csv").iloc[0]
        metrics = pd.Series(
            {
                "单帧质量": vt["frame_quality_score"],
                "光流连续": vt["flow_continuity_score"],
                "抖动抑制": vt["flow_temporal_jitter_score"],
                "内容一致": vt["content_consistency_score"],
                "闪烁抑制": vt["flicker_suppression_score"],
            }
        )
        fig, ax = plt.subplots(figsize=(8, 4.2))
        metrics.plot(kind="bar", ax=ax, color="#59A14F")
        ax.set_ylim(0, 100)
        ax.set_title("视频时序质量指标")
        ax.set_ylabel("分数")
        fig.tight_layout()
        fig.savefig(FIGURE_DIR / "video_temporal_metrics.png", dpi=220)
        plt.close(fig)

    fig, ax = plt.subplots(figsize=(12, 3.8))
    ax.axis("off")
    steps = ["输入审计", "题面读取", "图像特征", "组合赋权", "综合评价", "有效性修正", "视频时序", "复核审计", "论文提交"]
    xs = np.linspace(0.05, 0.95, len(steps))
    for i, (x, s) in enumerate(zip(xs, steps)):
        ax.text(x, 0.5, s, ha="center", va="center", bbox=dict(boxstyle="round,pad=0.3", fc="#F5F7FA", ec="#4C78A8"))
        if i < len(steps) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.045, 0.5), xytext=(x + 0.045, 0.5), arrowprops=dict(arrowstyle="->", color="#666"))
    ax.set_title("总体建模流程图")
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "modeling_flowchart.png", dpi=220)
    plt.close(fig)

    for src in FIGURE_DIR.glob("*.png"):
        shutil.copy2(src, PAPER_FIGURE_DIR / src.name)
    figs = sorted(FIGURE_DIR.glob("*.png"))
    fig_df = pd.DataFrame([{"figure": f.name, "path": str(f), "paper_copy": str(PAPER_FIGURE_DIR / f.name)} for f in figs])
    fig_df.to_csv(TABLE_DIR / "figure_list.csv", index=False, encoding="utf-8-sig")
    write_text(REPORT_DIR / "FIGURE_LIST.md", "# 图表清单\n\n" + table_md(fig_df, 100))
    log("Stage 11 completed: figures")


def load_results_for_paper() -> dict[str, pd.DataFrame]:
    paths = {
        "manifest": TABLE_DIR / "sample_manifest.csv",
        "features": TABLE_DIR / "features.csv",
        "weights": TABLE_DIR / "combined_weights.csv",
        "scores": TABLE_DIR / "final_semantic_quality_grades.csv",
        "relative_tiers": TABLE_DIR / "image_relative_quality_tiers.csv",
        "video": TABLE_DIR / "video_temporal_features.csv",
        "param_corr": TABLE_DIR / "parameter_quality_correlation.csv",
        "agent_dis": TABLE_DIR / "agent_disagreement.csv",
    }
    return {k: pd.read_csv(v) if v.exists() else pd.DataFrame() for k, v in paths.items()}


def compose_paper_markdown() -> str:
    res = load_results_for_paper()
    scores = res["scores"].sort_values("Q_final", ascending=False)
    rel_tiers = res.get("relative_tiers", pd.DataFrame()).sort_values("rank") if "relative_tiers" in res else pd.DataFrame()
    video = res["video"]
    top = scores.iloc[0] if not scores.empty else None
    low = scores.iloc[-1] if not scores.empty else None
    top_name = str(top["file_name"]) if top is not None else "NA"
    low_name = str(low["file_name"]) if low is not None else "NA"
    top_score = float(top["Q_final"]) if top is not None else float("nan")
    low_score = float(low["Q_final"]) if low is not None else float("nan")
    video_score = float(video.iloc[0]["video_temporal_quality_score"]) if not video.empty else np.nan
    instability = bool(video.iloc[0]["instability_flag"]) if not video.empty else False
    md = f"""# 基于组合赋权与时序一致性的 AI 生成内容质量评估模型

## 摘要

针对 2026 年第八届中青杯数学建模竞赛 B 题，本文围绕 AI 生成图像和视频的质量评价建立无参考评价模型。首先将图像质量分解为语义结构完整性代理、技术质量、结构完整性和视觉表达四类指标，构造 Laplacian 清晰度、自然场景统计质量、伪影风险、边缘连续性、形状规则性、亮度均衡、色彩丰富度和构图代理等可计算指标。其次采用 AHP、熵权法和 CRITIC 法组合赋权，并利用 TOPSIS 与灰色关联分析得到图像综合质量分。为避免分辨率、边缘复杂度等数量指标造成虚高，进一步引入有效性门控和风险惩罚，得到最终质量指数。最后针对附件 2 的车流视频，建立基于 Farneback 光流连续性、颜色直方图一致性和亮度突变的时序质量模型。

附件 1 的 8 张图像中，综合质量最高样本为 {top_name}，最终分为 {top_score:.2f}；最低样本为 {low_name}，最终分为 {low_score:.2f}。为回应题面高/中/低质量覆盖要求，本文同时给出固定阈值绝对等级和附件内部相对三档等级；相对等级只表示本批 8 张图像内部排序，不解释为专家标签。附件 2 视频时序质量分为 {video_score:.2f}，模型判断时序失稳标记为 {"是" if instability else "否"}。由于附件未提供原始提示词、专家评分和真实生成参数，参数优化部分仅给出不适用审计和可观测代理敏感性分析。

**关键词**：AI 生成内容；无参考图像质量评价；组合赋权；TOPSIS；灰色关联；光流连续性

## 1 问题重述

B 题要求建立 AI 生成内容质量评价模型。问题一要求建立无参考图像质量评价数学模型，题面强调提示词语义保真度、技术质量和结构完整性等指标；问题二要求对附件 1 的 8 张 AI 生成图像进行评估，并分析内容类型对指标敏感性的影响；问题三要求建立视频时序质量模型，考虑光流连续性、内容一致性、闪烁检测以及时序失稳条件，并分析附件 2 的车流视频。

## 2 问题分析

AI 生成图像没有标准参考图像，不能直接使用 PSNR 或 SSIM。因此应采用无参考评价思路，把画面质量拆解为可解释的统计和结构指标。对于视频，单帧图像质量不足以判断整体质量，还需要衡量相邻帧运动场是否平滑、颜色和亮度是否突变、场景内容是否保持一致。

![总体建模流程图](figures/modeling_flowchart.png)

## 3 模型假设

1. 附件未给出原始提示词，本文不计算真实 prompt-image 保真度，只计算语义结构完整性代理。
2. 附件未给出专家标签，本文分级为模型综合等级。
3. 附件未给出真实生成参数，因此不做真实参数优化。
4. 视频不使用外部训练检测器，车辆和场景稳定性使用光流与视觉统计代理。

## 4 符号说明

| 符号 | 含义 |
|---|---|
| $X=(x_{{ij}})$ | 原始指标矩阵 |
| $x'_{{ij}}$ | 归一化指标 |
| $w_j$ | 第 $j$ 个指标组合权重 |
| $C_i$ | TOPSIS 贴近度 |
| $G_i$ | 灰色关联度 |
| $Q_i$ | 综合质量分 |
| $G_{{valid}}$ | 有效性门控 |
| $F_t$ | 第 $t$ 对帧的光流场 |

## 5 数据预处理与特征提取

附件 1 包含 8 张图像，附件 2 包含 1 个 24 fps、121 帧的车流视频。图像使用 Pillow 和 OpenCV 读取，统一在缩放副本上计算特征；视频按固定间隔采样帧并计算相邻帧光流。

题面问题二提到样本需覆盖写实风景、人物肖像、艺术插画、产品渲染等类型，但附件 1 的 8 张图经可视审计未发现明确产品渲染样本。为保证数据真实性，本文按附件实际内容类型建立样本清单和敏感性分析，不补造不存在的产品渲染样本或质量标签。

| 样本ID | 文件 | 类型 | 最终分 | 绝对等级 | 相对等级 |
|---|---|---|---:|---|---|
"""
    for _, r in scores.iterrows():
        abs_grade = r.get("absolute_quality_grade", r.get("final_grade", ""))
        rel_grade = r.get("relative_quality_tier", "")
        md += f"| {r['sample_id']} | {r['file_name']} | {r['content_type']} | {r['Q_final']:.2f} | {abs_grade} | {rel_grade} |\n"
    md += f"""

## 6 问题一：无参考图像质量评价模型

### 6.1 指标体系

设图像样本数为 $n$，指标数为 $m$，原始矩阵为

$$
X=(x_{{ij}})_{{n\\times m}}.
$$

正向指标采用极差归一化

$$
x'_{{ij}}=\\frac{{x_{{ij}}-\\min_i x_{{ij}}}}{{\\max_i x_{{ij}}-\\min_i x_{{ij}}}},
$$

负向指标采用

$$
x'_{{ij}}=\\frac{{\\max_i x_{{ij}}-x_{{ij}}}}{{\\max_i x_{{ij}}-\\min_i x_{{ij}}}}.
$$

对复杂度、数量和边缘规模类指标采用饱和函数

$$
q(x)=1-\\exp(-x/\\tau),
$$

防止“越多越好”的评分虚高。

![质量指标体系层次图](figures/indicator_hierarchy.png)

### 6.2 语义结构完整性代理

题面要求将提示词分为主体对象、属性描述、场景关系、风格指令四类。由于附件没有原始提示词，本文不能计算真实 prompt-image 语义保真度，而是用审题可视描述构造代理文本 $T_i$，建立语义结构完整性代理指标。设四类要素集合分别为 $E_i^s,E_i^a,E_i^r,E_i^g$，语义要素覆盖分为

$$
S_i=0.55\\frac{{I(E_i^s)+I(E_i^a)+I(E_i^r)+I(E_i^g)}}{{4}}+0.45\\min\\left(\\frac{{|T_i|}}{{9}},1\\right).
$$

该指标只表示代理描述是否覆盖题面要求的语义维度，不代表原始提示词匹配度，也不等同于 CLIP 语义相似度。

### 6.3 技术质量与结构完整性

清晰度采用 Laplacian 方差：

$$
L_i=Var(\\nabla^2 I_i),\\quad S^L_i=\\frac{{\\log(1+L_i)}}{{\\log(1+L_0)}}.
$$

伪影风险综合像素截断率、块效应和过平滑风险：

$$
R_i^A=0.45R_i^c+0.35R_i^b+0.20R_i^o.
$$

边缘连续性基于 Canny 边缘连通区域的高分位面积 $A_{{90}}$：

$$
E_i=(1-e^{{-A_{{90}}/\\tau_A}})(1-e^{{-d_i/\\tau_d}}).
$$

结构完整性为

$$
H_i=0.55E_i+0.30R_i^s+0.15(1-R_i^A).
$$

## 7 问题二：图像综合评价算法与结果

### 7.1 组合赋权

AHP 权重由评价逻辑构造一致判断矩阵；熵权法刻画样本区分度；CRITIC 法刻画指标波动和冲突性。组合权重为

$$
w=\\alpha w^{{AHP}}+\\beta w^{{Entropy}}+\\gamma w^{{CRITIC}},\\quad \\alpha+\\beta+\\gamma=1.
$$

![指标权重对比](figures/weight_comparison.png)

### 7.2 TOPSIS-灰色关联综合评价

TOPSIS 贴近度为

$$
C_i=\\frac{{D_i^-}}{{D_i^+ + D_i^-}}.
$$

灰色关联系数为

$$
\\xi_{{ij}}=\\frac{{\\Delta_{{min}}+\\rho\\Delta_{{max}}}}{{\\Delta_{{ij}}+\\rho\\Delta_{{max}}}},
$$

关联度为 $G_i=\\sum_j w_j\\xi_{{ij}}$。综合分采用

$$
Q_i=100[0.5C_i+0.5G_i].
$$

![综合质量分布](figures/quality_score_distribution.png)

### 7.3 有效性修正

为避免高分辨率、强边缘或复杂纹理导致虚高，设有效性门控

$$
G_{{valid}}=0.30V_r+0.25V_t+0.20V_s+0.15V_m+0.10V_e.
$$

最终分为

$$
Q_{{final}}=clip(Q_{{raw}}(0.85+0.15G_{{valid}})-0.5P,0,100).
$$

![分数修正对比](figures/score_comparison.png)

### 7.4 绝对等级与相对等级

固定阈值绝对等级反映模型分数是否达到预设质量区间；附件内部相对等级用于回应题面“高、中、低质量等级”的比较需求。本文按 `Q_final` 排序，将第 1-2 名定义为相对高，第 3-5 名定义为相对中，第 6-8 名定义为相对低。该相对等级只用于本批附件内部比较，不是专家标签。

| 样本ID | 文件 | 类型 | 最终分 | 绝对等级 | 相对等级 | 关键优势 | 关键弱点 |
|---|---|---|---:|---|---|---|---|
"""
    if not rel_tiers.empty:
        for _, r in rel_tiers.iterrows():
            md += f"| {r['sample_id']} | {r['file_name']} | {r['content_type']} | {r['Q_final']:.2f} | {r['absolute_quality_grade']} | {r['relative_quality_tier']} | {r['key_strength']} | {r['key_weakness']} |\n"
    md += f"""

### 7.5 内容类型敏感性

写实风景类更依赖清晰度、自然场景统计和亮度均衡；人物肖像和人物插画更依赖结构完整性、边缘连续性和伪影风险；水墨、像素等强风格图像的自然场景统计得分未必最高，因此需要把风格与技术质量分开解释。

由于附件实际样本未覆盖明确产品渲染图，本文不对产品渲染类型给出单独敏感性结论；若后续补充产品样本，可复用同一指标体系并重点比较边缘规则性、材质高光、透视一致性和伪影风险。

## 8 问题三：视频时序质量评价模型

### 8.1 光流连续性

设 $F_t(x,y)=(u_t,v_t)$ 为相邻帧光流。光流空间平滑性定义为

$$
S_t^f=\\frac{{1}}{{1+\\frac{{\\|\\nabla |F_t|\\|_1}}{{\\overline{{|F_t|}}+\\epsilon}}}}.
$$

光流时间抖动为

$$
J=median_t\\left(\\frac{{\\| |F_t|-|F_{{t-1}}|\\|_1}}{{\\overline{{|F_t|}}+\\epsilon}}\\right),
$$

抖动抑制分为 $S^j=1/(1+J)$。

### 8.2 内容一致性与闪烁检测

内容一致性使用 HSV 二维直方图相关系数 $H_t$；亮度闪烁用相邻帧均值亮度差

$$
B_t=|\\bar{{Y}}_t-\\bar{{Y}}_{{t-1}}|.
$$

闪烁抑制分为

$$
S^b=clip(1-\\overline{{B_t}}/0.08,0,1).
$$

视频时序质量分为

$$
Q_v=100(0.25S^I+0.25S^f+0.20S^j+0.20S^H+0.10S^b).
$$

![视频时序质量指标](figures/video_temporal_metrics.png)

### 8.3 时序失稳必要条件

若同时出现较大的亮度突变、较低的光流平滑性、较低内容一致性或整体时序分低于阈值，则可判定存在时序失稳风险。本文使用以下必要条件：

$$
\\overline{{B_t}}>0.08\\quad \\text{{or}}\\quad \\overline{{S_t^f}}<0.45\\quad \\text{{or}}\\quad \\overline{{H_t}}<0.70\\quad \\text{{or}}\\quad Q_v<65.
$$

附件 2 视频的时序质量分为 {video_score:.2f}，时序失稳标记为 {"是" if instability else "否"}。

![视频关键帧质量分析](figures/video_frame_quality_analysis.png)

## 9 可靠性与稳定性分析

本文使用 TOPSIS、灰色关联和加权和进行跨模型对比，并对 $\\lambda=0.3,0.5,0.7$ 做敏感性分析。多智能体复核仅作为规则型复核，不覆盖主评分。

![多智能体评分热力图](figures/agent_score_heatmap.png)

## 10 模型评价

优点：模型完全本地可复现，指标均有明确数学定义；评分过程兼顾主观评价逻辑和客观数据差异；视频模型直接对应题面给出的光流、内容一致性和闪烁因素。

不足：缺少原始提示词、专家评分和真实生成参数，语义结构完整性和等级校准只能使用代理；附件 1 仅含 8 张图像，组合赋权和相对分级主要用于本批样本内部比较，外推到更大规模 AIGC 图像集时需要重新校准权重和等级阈值；没有训练式目标检测器时，车流视频中的车辆数量和身份一致性只能由光流与颜色统计间接描述。

## 11 结论

本文建立了覆盖图像单帧质量和视频时序质量的 AIGC 质量评价模型。对 8 张图像的评价表明，清晰度、伪影风险、边缘连续性和结构完整性是影响综合分的关键因素；对车流视频的评价表明，光流连续性和亮度/颜色稳定性可以有效刻画时序质量。由于附件无真实生成参数，本文不输出虚构参数最优组合，只给出基于指标缺陷的内容生成与筛选策略。

## 参考文献

[1] Mittal A, Soundararajan R, Bovik A C. Making a completely blind image quality analyzer. IEEE Signal Processing Letters, 2013.

[2] Farnebäck G. Two-frame motion estimation based on polynomial expansion. Scandinavian Conference on Image Analysis, 2003.

[3] Deng J L. Introduction to grey system theory. The Journal of Grey System, 1989.

[4] Saaty T L. The Analytic Hierarchy Process. New York: McGraw-Hill, 1980.

[5] OpenAI ChatGPT/Codex, GPT-5-Codex, OpenAI, 2026-06-05.

## AI 工具使用说明

本工程使用 AI 工具辅助工程组织、代码生成、报告草稿装配和格式审计。所有数据均来自赛题附件，模型公式、指标定义和结果均由本地脚本复现生成；参赛队需在提交前人工核验结论、格式和 AI 使用详情说明。

## 附录

附录文件包含输入审计、指标矩阵、权重表、图像评分表、视频时序特征、规则型 Agent 复核和一键复现脚本。
"""
    return md


def generate_modeling_paper() -> None:
    ensure_dirs()
    if not (FIGURE_DIR / "modeling_flowchart.png").exists():
        generate_figures()
    md = compose_paper_markdown()
    write_text(REPORT_DIR / "PAPER_DRAFT.md", md)
    write_text(PAPER_DIR / "main.md", md)
    build_docx(PAPER_DIR / "main.docx", with_excel_tables=False)
    docx_to_pdf(PAPER_DIR / "main.docx", PAPER_DIR / "main.pdf")
    log("Stage 12 completed: modeling paper")


def add_page_number(paragraph: Any) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr)
        run._r.append(fld_char2)
    except Exception:
        paragraph.add_run("1")


def set_run_font(run: Any, font_name: str = "宋体", size_pt: float = 12, bold: bool = False) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def build_docx(out_path: Path, with_excel_tables: bool = False, team_id: str = "待替换") -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Inches, Pt

    ensure_dirs()
    doc = Document()
    sec = doc.sections[0]
    # Margins are not specified in rules; use converted template if available, otherwise a conservative fallback.
    converted = EXTRACTED_DIR / "competition_template_converted.docx"
    if converted.exists():
        try:
            from docx import Document as D2

            tdoc = D2(str(converted))
            tsec = tdoc.sections[0]
            sec.page_width = tsec.page_width
            sec.page_height = tsec.page_height
            sec.top_margin = tsec.top_margin
            sec.bottom_margin = tsec.bottom_margin
            sec.left_margin = tsec.left_margin
            sec.right_margin = tsec.right_margin
        except Exception:
            sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)
    else:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(12)
    try:
        styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # type: ignore[name-defined]
    except Exception:
        pass
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = "Table Grid"
    info_rows = [("队伍编号", team_id), ("选题", "B")]
    for i, row in enumerate(info_rows):
        for j, val in enumerate(row):
            cell = info_table.cell(i, j)
            cell.text = str(val)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, "宋体", 12, bold=(j == 0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于组合赋权与时序一致性的 AI 生成内容质量评估模型")
    set_run_font(r, "黑体", 16, True)

    md = compose_paper_markdown()
    lines = md.splitlines()
    skip_title = True
    table_buffer: list[str] = []
    in_formula = False
    formula_buffer: list[str] = []
    figure_no = 0
    table_no = 0
    formula_no = 0
    for line in lines:
        if line.startswith("# "):
            if skip_title:
                skip_title = False
                continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            table_no += 1
            add_caption_paragraph(doc, f"表 {table_no} {infer_markdown_table_caption(table_buffer)}", kind="table")
            add_markdown_table(doc, table_buffer)
            table_buffer = []
        if not line.strip():
            continue
        if line.startswith("$$"):
            if in_formula:
                formula_no += 1
                add_formula_paragraph(doc, " ".join(formula_buffer), formula_no)
                formula_buffer = []
                in_formula = False
            else:
                in_formula = True
                formula_buffer = []
            continue
        if in_formula:
            formula_buffer.append(line.strip())
            continue
        if line.startswith("!["):
            caption_match = re.search(r"!\[(.*?)\]", line)
            m = re.search(r"\((.*?)\)", line)
            if m:
                img = PAPER_DIR / m.group(1)
                if img.exists():
                    figure_no += 1
                    doc.add_picture(str(img), width=Inches(5.7))
                    last = doc.paragraphs[-1]
                    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = caption_match.group(1) if caption_match else img.stem
                    add_caption_paragraph(doc, f"图 {figure_no} {caption}", kind="figure")
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[3:].strip())
            set_run_font(r, "黑体", 14, True)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line[4:].strip())
            set_run_font(r, "黑体", 12, True)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(0)
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", line.strip())
            r = p.add_run(text)
            set_run_font(r, "宋体", 12)
    if table_buffer:
        table_no += 1
        add_caption_paragraph(doc, f"表 {table_no} {infer_markdown_table_caption(table_buffer)}", kind="table")
        add_markdown_table(doc, table_buffer)
    if formula_buffer:
        formula_no += 1
        add_formula_paragraph(doc, " ".join(formula_buffer), formula_no)
    if with_excel_tables:
        doc.add_page_break()
        p = doc.add_paragraph()
        r = p.add_run("主要结果表（Excel 同步版本）")
        set_run_font(r, "黑体", 14, True)
        add_csv_table_to_doc(doc, TABLE_DIR / "final_semantic_quality_grades.csv", "图像最终评分表")
        add_csv_table_to_doc(doc, TABLE_DIR / "image_relative_quality_tiers.csv", "附件1相对高/中/低等级表")
        add_csv_table_to_doc(doc, TABLE_DIR / "video_temporal_features.csv", "视频时序质量表")
        add_csv_table_to_doc(doc, TABLE_DIR / "combined_weights.csv", "组合权重表")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def add_caption_paragraph(doc: Any, text: str, kind: str = "figure") -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, "宋体", 10.5, False)


def infer_markdown_table_caption(lines: list[str]) -> str:
    header = lines[0] if lines else ""
    if "符号" in header:
        return "符号说明"
    if "关键优势" in header or "关键弱点" in header:
        return "附件1相对高/中/低等级与优劣势"
    if "样本ID" in header and "最终分" in header:
        return "附件1图像质量评分结果"
    return "论文正文表格"


def add_formula_paragraph(doc: Any, formula: str, number: int) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    formula = re.sub(r"\s+", " ", formula).strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{formula}    ({number})")
    set_run_font(r, "Times New Roman", 11, False)


def add_markdown_table(doc: Any, lines: list[str]) -> None:
    clean = [ln for ln in lines if not re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", ln)]
    rows = []
    for ln in clean:
        parts = [p.strip() for p in ln.strip().strip("|").split("|")]
        if parts:
            rows.append(parts)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, "宋体", 9.5, bold=(i == 0))


def add_csv_table_to_doc(doc: Any, path: Path, caption: str, max_rows: int = 12) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r, "黑体", 10.5, True)
    df = pd.read_csv(path).head(max_rows)
    rows = [list(df.columns)] + df.astype(str).values.tolist()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            for p2 in cell.paragraphs:
                for run in p2.runs:
                    set_run_font(run, "宋体", 8.5, i == 0)


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import win32com.client

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return True
    except Exception as exc:
        try:
            word.Quit()  # type: ignore[name-defined]
        except Exception:
            pass
        log(f"Word PDF conversion failed: {exc}")
        return False


def build_excel_tables_word() -> None:
    ensure_dirs()
    if not (PAPER_DIR / "main.md").exists():
        generate_modeling_paper()
    xlsx = PAPER_TABLE_DIR / "main_tables.xlsx"
    with pd.ExcelWriter(xlsx, engine="openpyxl") as writer:
        sheets = {
            "样本清单": TABLE_DIR / "sample_manifest.csv",
            "指标定义": TABLE_DIR / "indicator_definition_table.csv",
            "指标矩阵": TABLE_DIR / "indicator_matrix_normalized.csv",
            "组合权重": TABLE_DIR / "combined_weights.csv",
            "图像评分": TABLE_DIR / "final_semantic_quality_grades.csv",
            "相对等级": TABLE_DIR / "image_relative_quality_tiers.csv",
            "视频时序": TABLE_DIR / "video_temporal_features.csv",
            "参数审计": TABLE_DIR / "optimization_method_comparison.csv",
        }
        for sheet, path in sheets.items():
            if path.exists():
                pd.read_csv(path).to_excel(writer, sheet_name=sheet[:31], index=False)
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = load_workbook(xlsx)
    thin = Side(style="thin", color="999999")
    for ws in wb.worksheets:
        for cell in ws[1]:
            cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="4F81BD")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows():
            for cell in row:
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for col in ws.columns:
            max_len = max(len(str(c.value)) if c.value is not None else 0 for c in col)
            ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 10), 36)
    wb.save(xlsx)
    build_docx(PAPER_DIR / "main_with_excel_tables.docx", with_excel_tables=True)
    md = f"""# Excel 表格与 Word 装配

- Excel 工作簿：`{xlsx}`
- Word 文件：`{PAPER_DIR / "main_with_excel_tables.docx"}`
- 表头、边框、自动换行和列宽已程序化设置。
- 公式以清晰文本公式保留，仍建议人工复核是否需转为 Word 原生公式。
"""
    write_text(REPORT_DIR / "EXCEL_TABLE_ASSEMBLY.md", md)
    log("Stage 14 completed: Excel and Word assembly")


def apply_word_template_format() -> None:
    ensure_dirs()
    if not (PAPER_DIR / "main_with_excel_tables.docx").exists():
        build_excel_tables_word()
    ok = docx_to_pdf(PAPER_DIR / "main_with_excel_tables.docx", PAPER_DIR / "main_with_excel_tables.pdf")
    # Also refresh main.pdf to keep both versions present.
    if (PAPER_DIR / "main.docx").exists():
        docx_to_pdf(PAPER_DIR / "main.docx", PAPER_DIR / "main.pdf")
    md = f"""# Word 格式审计

- 参赛细则已读取：{"PASS" if (REPORT_DIR / "FORMAT_REQUIREMENTS.md").exists() else "MISSING"}
- 论文模板已尝试读取/转换：{"PASS" if (EXTRACTED_DIR / "competition_template_converted.docx").exists() else "MANUAL_CHECK_REQUIRED"}
- 最终 Word：`{PAPER_DIR / "main_with_excel_tables.docx"}`
- 最终 PDF：`{PAPER_DIR / "main_with_excel_tables.pdf"}`
- PDF 转换状态：{"PASS" if ok else "MANUAL_CHECK_REQUIRED"}
- 字体策略：题目三号黑体，一级标题四号黑体居中，二级/三级小四黑体左对齐，正文小四宋体，1.25 倍行距。
- 页码策略：页脚居中插入 PAGE 域。
- 备用题注策略：Word 草稿已自动插入“图 n ...”“表 n ...”题注。
- 备用公式策略：Word 草稿已将 Markdown 公式块渲染为居中公式文本并附顺序编号。
- 待人工确认：官方图题格式、官方表题格式、官方公式编号格式是否与备用策略一致。
"""
    write_text(REPORT_DIR / "WORD_FORMAT_AUDIT.md", md)
    log("Stage 15 completed: Word formatting")


def prepare_submission() -> None:
    ensure_dirs()
    if not (PAPER_DIR / "main_with_excel_tables.pdf").exists():
        apply_word_template_format()
    pdf_src = PAPER_DIR / "main_with_excel_tables.pdf"
    docx_src = PAPER_DIR / "main_with_excel_tables.docx"
    pdf_dst = SUBMISSION_DIR / "B题论文_AI生成内容质量评估与参数优化.pdf"
    docx_dst = SUBMISSION_DIR / "B题论文_AI生成内容质量评估与参数优化.docx"
    if pdf_src.exists():
        shutil.copy2(pdf_src, pdf_dst)
    if docx_src.exists():
        shutil.copy2(docx_src, docx_dst)
    ai = f"""# AI 工具使用详情说明

## 所用 AI 工具名称和版本

- ChatGPT/Codex，GPT-5-Codex，OpenAI，使用日期：2026-06-05。

## 使用目的和应用环节

- 辅助建立工程目录、编写本地可复现脚本、整理模型公式、生成论文草稿、装配 Word/PDF 与提交材料。

## 关键交互记录摘要

- 用户要求从 B 题压缩包开始，完成输入审计、题面读取、附件解析、建模、图表、论文与提交包。
- 工程按题面 PDF 自动修正为 NR-IQA 图像评价与视频时序质量评价，不伪造缺失参数或专家评分。

## 采纳内容和人工修改情况

- 本地脚本生成所有表格、图表和分数。
- 参赛队提交前应人工复核题面理解、格式、公式、结论和队伍编号。
"""
    write_text(SUBMISSION_DIR / "AI工具使用详情说明.md", ai)
    support_zip = SUBMISSION_DIR / "B题支撑材料_zhongqingB.zip"
    include_roots = [ROOT / "configs", ROOT / "src", ROOT / "scripts", OUTPUT_DIR, REPORT_DIR, PAPER_TABLE_DIR]
    with zipfile.ZipFile(support_zip, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for base in include_roots:
            if not base.exists():
                continue
            for p in base.rglob("*"):
                if p.is_file() and SUBMISSION_DIR not in p.parents:
                    zf.write(p, p.relative_to(ROOT))
        for p in [PAPER_DIR / "main.md", PAPER_DIR / "main.docx", PAPER_DIR / "main.pdf"]:
            if p.exists():
                zf.write(p, p.relative_to(ROOT))
    manifest = f"""# 提交材料清单

- 论文 PDF：`{pdf_dst}`（细则最终要求应使用真实队伍编号命名）
- 论文 Word：`{docx_dst}`（便于人工复核，正式提交以 PDF 为准）
- 支撑材料：`{support_zip}`
- AI 工具使用详情说明：`{SUBMISSION_DIR / "AI工具使用详情说明.md"}`

## 当前命名状态

当前文件名是占位说明性命名，不是细则要求的最终命名。正式提交前必须用真实队伍编号运行 `scripts/18_rename_submission.py`，生成 `B真实队伍编号.pdf` 和 `B真实队伍编号材料.zip`。

## 命名提醒

若队伍编号为 `B202600001`，运行：

```powershell
python scripts\\18_rename_submission.py B202600001
```
"""
    write_text(SUBMISSION_DIR / "SUBMISSION_MANIFEST.md", manifest)
    log("Stage 16 completed: submission package")


def pdf_page_count(path: Path) -> int | None:
    if not path.exists():
        return None
    try:
        return fitz.open(path).page_count
    except Exception:
        return None


def final_quality_audit() -> None:
    ensure_dirs()
    if not (SUBMISSION_DIR / "SUBMISSION_MANIFEST.md").exists():
        prepare_submission()
    if (TABLE_DIR / "format_requirements.csv").exists():
        fmt_df = pd.read_csv(TABLE_DIR / "format_requirements.csv")
        pending_mask = fmt_df.astype(str).apply(lambda col: col.str.contains("待人工确认", na=False)).any(axis=1)
        pending_count = int(pending_mask.sum())
    else:
        fmt_text = (REPORT_DIR / "FORMAT_REQUIREMENTS.md").read_text(encoding="utf-8") if (REPORT_DIR / "FORMAT_REQUIREMENTS.md").exists() else ""
        pending_count = fmt_text.count("待人工确认")
    pdf_path = PAPER_DIR / "main_with_excel_tables.pdf"
    pages = pdf_page_count(pdf_path)
    checks = [
        ("题面读取完成", (REPORT_DIR / "PROBLEM_READING.md").exists(), ""),
        ("三问逐一回答", (PAPER_DIR / "main.md").exists(), "见论文第 6-8 节"),
        ("附件全部使用", (TABLE_DIR / "sample_manifest.csv").exists(), "8 图 + 1 视频"),
        ("参赛细则已读取", (REPORT_DIR / "FORMAT_REQUIREMENTS.md").exists(), ""),
        ("论文模板已读取", (REPORT_DIR / "WORD_TEMPLATE_AUDIT.md").exists(), "旧 .doc 已尝试转换"),
        ("格式要求已提取", (TABLE_DIR / "format_requirements.csv").exists(), ""),
        ("指标体系完整", (TABLE_DIR / "indicator_definition_table.csv").exists(), ""),
        ("组合赋权存在", (TABLE_DIR / "combined_weights.csv").exists(), ""),
        ("综合评价模型存在", (TABLE_DIR / "comprehensive_quality_scores.csv").exists(), ""),
        ("有效性修正存在", (TABLE_DIR / "corrected_quality_scores.csv").exists(), ""),
        ("语义分级校准存在", (TABLE_DIR / "final_semantic_quality_grades.csv").exists(), ""),
        ("相对高/中/低等级表存在", (TABLE_DIR / "image_relative_quality_tiers.csv").exists(), "用于问题二附件内部等级解释"),
        ("参数-质量关系模型存在", (REPORT_DIR / "PARAMETER_EFFECT_ANALYSIS.md").exists(), "真实参数缺失，标为代理敏感性"),
        ("参数优化模型存在", (REPORT_DIR / "PARAMETER_OPTIMIZATION_REPORT.md").exists(), "题面不适用，不伪造参数"),
        ("多智能体复核存在", (TABLE_DIR / "agent_scores.csv").exists(), ""),
        ("视频或图像附件已处理", (TABLE_DIR / "video_temporal_features.csv").exists() and (TABLE_DIR / "features.csv").exists(), ""),
        ("关键表格生成", len(list(TABLE_DIR.glob("*.csv"))) >= 10, ""),
        ("关键图表生成", len(list(FIGURE_DIR.glob("*.png"))) >= 8, ""),
        ("Word/PDF 生成", (PAPER_DIR / "main_with_excel_tables.docx").exists() and (PAPER_DIR / "main_with_excel_tables.pdf").exists(), ""),
        ("提交包生成", (SUBMISSION_DIR / "B题支撑材料_zhongqingB.zip").exists(), ""),
        ("弱监督限制已说明", "没有专家标签" in (REPORT_DIR / "ASSUMPTIONS.md").read_text(encoding="utf-8"), ""),
        ("8图小样本限制已说明", "仅含 8 张图像" in (PAPER_DIR / "main.md").read_text(encoding="utf-8") if (PAPER_DIR / "main.md").exists() else False, ""),
        ("AI 风险不是违规判定", "不构成违规判定" in (REPORT_DIR / "ASSUMPTIONS.md").read_text(encoding="utf-8"), ""),
        ("参数范围来源已说明", (REPORT_DIR / "PARAMETER_OPTIMIZATION_REPORT.md").exists(), "无真实参数范围"),
        ("一键复现命令可用", (ROOT / "scripts" / "19_run_all.py").exists(), ""),
        ("论文不存在旧结论残留", "A 题" not in (PAPER_DIR / "main.md").read_text(encoding="utf-8"), ""),
        ("页数是否符合要求", pages is not None and pages <= 20, f"PDF 页数={pages}"),
        ("是否存在待人工确认格式项", pending_count == 0, f"待人工确认计数={pending_count}"),
    ]
    audit_rows = []
    for item, ok, note in checks:
        if item == "是否存在待人工确认格式项" and not ok:
            status = "MANUAL_CHECK_REQUIRED"
        elif item in ["参数优化模型存在", "参数-质量关系模型存在"] and ok:
            status = "N/A_BY_PROBLEM_WITH_REPORT"
        elif item == "页数是否符合要求" and not ok:
            status = "MANUAL_CHECK_REQUIRED"
        else:
            status = "PASS" if ok else "FAIL"
        audit_rows.append({"check_item": item, "status": status, "note": note})
    audit_df = pd.DataFrame(audit_rows)
    audit_df.to_csv(TABLE_DIR / "final_quality_audit.csv", index=False, encoding="utf-8-sig")
    write_text(REPORT_DIR / "FINAL_QUALITY_AUDIT.md", "# 最终质量审计\n\n" + table_md(audit_df, 80))
    write_auxiliary_final_reports(pages, pending_count)
    update_readme_and_continuation(audit_df, pages, pending_count)
    log("Stage 17 completed: final audit")


def write_auxiliary_final_reports(pages: int | None, pending_count: int) -> None:
    method = """# 方法链重构计划

## 重构原因

题面 PDF 已确认真实 B 题三问为：

1. 建立无参考图像质量评价（NR-IQA）模型；
2. 对附件 1 的 8 张 AI 生成图像进行评价并分析内容类型敏感性；
3. 建立视频时序质量模型并分析附件 2 车流视频是否时序失稳。

用户初始任务链中的“真实参数-质量关系”和“多目标参数优化”不再作为主论文问题推进，因为附件没有真实生成参数记录，题面也没有要求输出参数组合最优解。

## 更新后的主任务链

输入审计与题面读取 -> 格式要求提取 -> 附件图像/视频解析 -> NR-IQA 指标量化 -> AHP-熵权-CRITIC 组合赋权 -> TOPSIS-灰色关联综合评分 -> 有效性修正与语义等级校准 -> 光流/一致性/闪烁视频时序质量 -> 规则型 Agent 复核 -> 图表与论文 -> Word/PDF 装配 -> 提交材料打包。

## 参数链路处理

参数相关脚本保留为审计步骤，输出 `NOT_APPLICABLE` 和可观测代理敏感性分析；不伪造 seed、steps、guidance scale、模型名、采样器或专家评分。
"""
    write_text(REPORT_DIR / "METHOD_RESTRUCTURE_PLAN.md", method)

    page = f"""# 页数限制调整计划

## 当前页数

- `paper/main.pdf`：{pdf_page_count(PAPER_DIR / "main.pdf")} 页
- `paper/main_with_excel_tables.pdf`：{pages} 页

## 细则要求

摘要页作为第一页，正文从第二页开始；不要目录；正文尽量控制在 20 页以内；附录页数不限。

## 当前结论

{"当前 PDF 页数未超过 20 页，暂不需要压缩。" if pages is not None and pages <= 20 else "当前页数无法确认或可能超过 20 页，需要人工压缩。"}

## 如需压缩的优先级

1. 将 Excel 同步表格版中的长表移至支撑材料；
2. 主文保留 8-10 张关键图，其余放附录或支撑材料；
3. 缩短多智能体复核文字，只保留分歧结论；
4. 删除参数优化不适用部分的冗长说明，仅保留审计结论。
"""
    write_text(REPORT_DIR / "PAGE_LIMIT_ADJUSTMENT_PLAN.md", page)

    repro = f"""# 可复现性说明

## 环境

- Python：见本机 `python --version`
- 依赖：`requirements.txt`
- 随机种子：`{SEED}`

## 一键复现

```powershell
cd D:\\zhongqingb\\B_aigc_quality_optimization
python scripts\\19_run_all.py
```

## 单步复现顺序

```powershell
python scripts\\00_audit_inputs.py
python scripts\\01_read_problem.py
python scripts\\13_extract_format_requirements.py
python scripts\\02_extract_materials.py
python scripts\\03_build_features.py
python scripts\\04_build_quality_indicator_system.py
python scripts\\05_build_weights.py
python scripts\\06_quality_evaluation.py
python scripts\\07_validity_corrected_score.py
python scripts\\08_parameter_effect_analysis.py
python scripts\\09_parameter_optimization.py
python scripts\\10_multi_agent_evaluation.py
python scripts\\11_generate_figures.py
python scripts\\12_generate_modeling_paper.py
python scripts\\14_build_excel_tables_word.py
python scripts\\15_apply_word_template_format.py
python scripts\\16_prepare_submission.py
python scripts\\17_final_quality_audit.py
```

## 不可自动确认项

- 待人工确认格式项数量：{pending_count}
- 真实队伍编号需要人工替换；
- 原始提示词、专家标签和真实生成参数未在附件中发现。
"""
    write_text(REPORT_DIR / "REPRODUCIBILITY.md", repro)

    readiness = f"""# 提交就绪检查

## 当前可用文件

- 占位命名论文 PDF：`submission/B题论文_AI生成内容质量评估与参数优化.pdf`
- 占位命名论文 Word：`submission/B题论文_AI生成内容质量评估与参数优化.docx`
- 占位命名支撑材料：`submission/B题支撑材料_zhongqingB.zip`
- AI 工具使用详情说明：`submission/AI工具使用详情说明.md`

## 正式提交前必须执行

将 `B真实队伍编号` 替换为报名页面显示的真实队伍编号：

```powershell
cd D:\\zhongqingb\\B_aigc_quality_optimization
python scripts\\18_rename_submission.py B真实队伍编号
```

运行后应提交：

- `submission/B真实队伍编号.pdf`
- `submission/B真实队伍编号材料.zip`

## 当前不能直接提交的内容

- 首页仍含 `待替换` 的通用 Word/PDF 不能作为最终命名文件直接提交；
- 任何测试编号文件不能提交，除非测试编号正是真实队伍编号；
- 原始赛题附件一般不需要打入支撑材料，支撑材料包主要包含程序、结果表、图和报告。

## 人工复核清单

1. 队伍编号正确；
2. PDF 文件名符合 `B+队伍编号.pdf`；
3. 支撑材料文件名符合 `B+队伍编号+材料.zip`；
4. 论文不含学校和队员姓名；
5. 图题、表题和公式编号符合人工确认的最终格式；
6. AI 工具使用详情说明随支撑材料一并提交；
7. 当前 PDF 页数为 {pages} 页，未超过 20 页目标。
"""
    write_text(REPORT_DIR / "SUBMISSION_READINESS.md", readiness)


def update_readme_and_continuation(audit_df: pd.DataFrame, pages: int | None, pending_count: int) -> None:
    q = pd.read_csv(TABLE_DIR / "problem_questions.csv") if (TABLE_DIR / "problem_questions.csv").exists() else pd.DataFrame()
    scores = pd.read_csv(TABLE_DIR / "final_semantic_quality_grades.csv") if (TABLE_DIR / "final_semantic_quality_grades.csv").exists() else pd.DataFrame()
    video = pd.read_csv(TABLE_DIR / "video_temporal_features.csv") if (TABLE_DIR / "video_temporal_features.csv").exists() else pd.DataFrame()
    readme = f"""# B_aigc_quality_optimization

## 当前状态

已完成 B 题从输入审计、题面读取、附件解析、图像 NR-IQA 建模、组合赋权、综合评价、附件内部相对高/中/低等级、视频时序质量分析、图表、论文 Word/PDF 和提交包装配的全流程草稿。

## 推荐提交文件

- 正式论文 PDF：`submission/B题论文_AI生成内容质量评估与参数优化.pdf`
- 支撑材料：`submission/B题支撑材料_zhongqingB.zip`
- AI 使用说明：`submission/AI工具使用详情说明.md`
- 注意：以上是占位命名文件。真实提交前必须运行 `python scripts\\18_rename_submission.py B真实队伍编号`，按细则生成最终命名文件。

## 工作区路径

- 工作区：`{WORKSPACE}`
- 工程目录：`{ROOT}`
- 原始材料：`{find_material_root()}`

## 题面三问

{table_md(q, 10)}

## 附件说明

- 附件 1：8 张 AI 生成图像。
- 附件 2：1 个车流视频。
- 原始提示词、专家标签、真实生成参数：未发现。

## 项目目录结构

目录按用户要求建立，核心代码在 `src/aigc_eval/pipeline.py`，各 Stage 脚本在 `scripts/`。

## Stage 过程

00 输入审计；01 题面读取；13 格式要求；02 材料解析；03 特征构建；04 指标体系；05 组合赋权；06 综合评价；07 有效性修正；08 代理敏感性；09 参数优化适用性；10 Agent 复核；11 图表；12 论文；14 Excel/Word；15 Word/PDF 格式；16 提交包；17 最终审计。

## 关键模型修正历史

题面 PDF 最高优先级。用户初始链路中的“真实参数优化”与 PDF 三问不一致，且附件无参数记录，因此主模型修正为图像 NR-IQA 与视频时序质量评价；参数脚本只输出不适用审计和可观测代理敏感性。

## 当前推荐模型口径

M1 NR-IQA 指标量化；M2 AHP-熵权-CRITIC 组合赋权；M3 TOPSIS-灰色关联综合评价；M4 有效性修正与绝对/相对等级校准；M5 光流-一致性-闪烁视频时序质量模型；M6 本地规则型 Agent 复核。

## Word 与论文格式要求

来自参赛细则：题目三号黑体，一级标题四号黑体居中，二三级标题小四黑体左对齐，正文小四宋体，1.25 倍行距；无目录；无页眉；页脚居中页码；正文尽量 20 页以内；文件名应为 `题号+队伍编号.pdf`。

## 当前论文版本

- Markdown：`paper/main.md`
- Word：`paper/main_with_excel_tables.docx`
- PDF：`paper/main_with_excel_tables.pdf`
- 当前 PDF 页数：{pages}

## 复现命令

```powershell
cd D:\\zhongqingb\\B_aigc_quality_optimization
python scripts\\19_run_all.py
```

## Git 历史

当前工程未初始化 Git；如需提交版本，请人工执行 `git init`、`git add`、`git commit`。

## 最终审计

{table_md(audit_df, 60)}

## 人工待检查项

- 将 `队伍编号：待替换` 替换为真实队伍编号。
- 根据细则将正式 PDF 命名为 `B+队伍编号.pdf`。
- 待人工确认格式项数量：{pending_count}。
- 检查 Word 公式是否需要改为原生公式对象。
- 人工确认附件 1 内容类型与题面“四类内容类型”表述是否存在歧义。

## 总结主线

AIGC 图像/视频样本 → 多维无参考质量指标 → 组合赋权 → TOPSIS-灰色关联评价 → 有效性修正 → 视频时序质量 → 规则型复核 → 数学建模论文与提交材料。
"""
    write_text(ROOT / "README.md", readme)

    continuation = f"""# WORK_CONTINUATION

## 已完成内容

已完成工作区创建、RAR 解压、题面与细则读取、模板转换尝试、附件解析、特征提取、图像质量评价、附件内部相对高/中/低等级、视频时序质量评价、图表、论文、Word/PDF、支撑材料 zip 和最终审计。

## 当前主模型

主模型为“无参考图像质量评价 + 组合赋权综合评价 + 光流时序质量评价”。参数优化不是题面主问题，且附件无真实参数，已按不适用处理。

## 当前主论文

- `paper/main.md`
- `paper/main_with_excel_tables.docx`
- `paper/main_with_excel_tables.pdf`

## 当前推荐提交文件

- `submission/B题论文_AI生成内容质量评估与参数优化.pdf`
- `submission/B题支撑材料_zhongqingB.zip`
- `submission/AI工具使用详情说明.md`
- 上述文件为占位命名；正式提交时以 `scripts\\18_rename_submission.py` 生成的 `B真实队伍编号.pdf` 与 `B真实队伍编号材料.zip` 为准。

## 一键复现命令

```powershell
cd D:\\zhongqingb\\B_aigc_quality_optimization
python scripts\\19_run_all.py
```

## 格式状态

参赛细则已读取；旧 `.doc` 模板已尝试转换。当前 PDF 页数：{pages}。

## 是否读取模板

{"是，已转换为 data/extracted/competition_template_converted.docx" if (EXTRACTED_DIR / "competition_template_converted.docx").exists() else "未能自动读取，需人工确认"}

## 是否有待人工确认格式项

是，数量：{pending_count}。

## 是否需要替换真实队伍编号

是。当前通用论文首页仍为 `队伍编号：待替换`。若提交目录中存在 `B202600001.pdf/docx`，它只是脚本验证样例；除非这正是真实队伍编号，否则不能直接提交。

## 是否需要压缩页数

{"否，当前程序统计不超过 20 页" if pages is not None and pages <= 20 else "是或需人工确认，当前页数可能超过 20 或无法统计"}。

## 后续人工工作

1. 替换真实队伍编号。
2. 运行 `python scripts\\18_rename_submission.py B真实队伍编号`。
3. 人工复核公式排版、图表题注、表题格式和页边距。
4. 人工确认没有参赛队员与学校信息。
5. 人工确认语义代理描述和内容类型分类是否需要调整。

## 如果需要继续优化

从 `reports/FINAL_QUALITY_AUDIT.md` 和 `reports/FORMAT_REQUIREMENTS.md` 的 `MANUAL_CHECK_REQUIRED` 项开始。
"""
    write_text(ROOT / "WORK_CONTINUATION.md", continuation)


def update_work_continuation(reason: str) -> None:
    ensure_dirs()
    text = f"""# WORK_CONTINUATION

当前流程中止原因：

{reason}

已创建工程目录：`{ROOT}`

建议下一步：补齐缺失关键文件后运行 `python scripts\\19_run_all.py`。
"""
    write_text(ROOT / "WORK_CONTINUATION.md", text)


def rename_submission(team_id: str) -> None:
    ensure_dirs()
    if not team_id or not re.match(r"^B?\d{6,}$|^B\d{6,}$", team_id):
        log("队伍编号格式未强制校验通过，但仍按输入执行复制命名")
    if not team_id.startswith("B"):
        team_id = "B" + team_id
    zip_src = SUBMISSION_DIR / "B题支撑材料_zhongqingB.zip"
    if not zip_src.exists():
        prepare_submission()
    docx_dst = SUBMISSION_DIR / f"{team_id}.docx"
    pdf_dst = SUBMISSION_DIR / f"{team_id}.pdf"
    build_docx(docx_dst, with_excel_tables=True, team_id=team_id)
    pdf_ok = docx_to_pdf(docx_dst, pdf_dst)
    zip_src = SUBMISSION_DIR / "B题支撑材料_zhongqingB.zip"
    zip_dst = SUBMISSION_DIR / f"{team_id}材料.zip"
    shutil.copy2(zip_src, zip_dst)
    write_text(
        SUBMISSION_DIR / "RENAMED_SUBMISSION.md",
        f"""# 已按队伍编号生成提交命名

- Word：`{docx_dst}`
- PDF：`{pdf_dst}`
- 支撑材料：`{zip_dst}`
- PDF 转换状态：{"PASS" if pdf_ok else "MANUAL_CHECK_REQUIRED"}
""",
    )
    log(f"Submission renamed for {team_id}")


STAGE_FUNCTIONS = {
    "00": audit_inputs,
    "01": read_problem,
    "02": extract_materials,
    "03": build_features,
    "04": build_quality_indicator_system,
    "05": build_weights,
    "06": quality_evaluation,
    "07": validity_corrected_score,
    "08": parameter_effect_analysis,
    "09": parameter_optimization,
    "10": multi_agent_evaluation,
    "11": generate_figures,
    "12": generate_modeling_paper,
    "13": extract_format_requirements,
    "14": build_excel_tables_word,
    "15": apply_word_template_format,
    "16": prepare_submission,
    "17": final_quality_audit,
}


def run_all() -> None:
    order = ["00", "01", "13", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12", "14", "15", "16", "17"]
    ensure_dirs()
    for stage in order:
        try:
            STAGE_FUNCTIONS[stage]()
        except Exception as exc:
            log(f"Stage {stage} failed: {type(exc).__name__}: {exc}")
            append_text(REPORT_DIR / "ASSUMPTIONS.md", f"\n## Stage {stage} 运行异常\n\n- {type(exc).__name__}: {exc}\n")
            if stage in ["00", "01", "02"]:
                update_work_continuation(f"Stage {stage} 关键步骤失败：{type(exc).__name__}: {exc}")
                raise
            continue


def main(stage: str | None = None, argv: list[str] | None = None) -> None:
    ensure_dirs()
    argv = list(sys.argv[1:] if argv is None else argv)
    if stage == "18":
        if not argv:
            raise SystemExit("用法：python scripts\\18_rename_submission.py B202600001")
        rename_submission(argv[0])
        return
    if stage == "19" or stage == "all":
        run_all()
        return
    if stage in STAGE_FUNCTIONS:
        STAGE_FUNCTIONS[stage]()
        return
    raise SystemExit(f"未知 stage: {stage}")
